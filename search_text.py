#!/usr/bin/env python3
"""
Usage:
    python search_text.py "text query" OCR_ROOT [threshold]

- text query = 2-12 word text to search for
- OCR_ROOT   = folderul cu subfolderele ce conțin *.txt din OCR
- threshold  = minimum score (0-100), default 70

Returns all matches above threshold with folder name and image name.
"""

import sys
import re
from pathlib import Path
from rapidfuzz import fuzz
from docx import Document

# Header de pagină în fișierele OCR (*.txt)
# ex: === PAGE 151 (pag151.jpg) ===
PAGE_HEADER_RE = re.compile(r"^=== PAGE (.+?) \((.+?)\) ===")

# Pagini cu foarte puține cuvinte sunt ignorate (zgomot)
MIN_PAGE_WORDS = 6

# Default threshold pentru matching
DEFAULT_THRESHOLD = 80


# ---------- UTILITĂȚI TEXT ----------

# Mapping pentru litere similare (pentru căutare OCR)
# Mapează diacritice la forma de bază pentru matching flexibil
# Notă: textul este deja lowercased înainte de aplicarea acestui mapping
SIMILAR_LETTERS_MAP = {
    # Diacritice românești
    'ș': 's',
    'ț': 't',
    'â': 'a',
    'î': 'a',  # î și â sunt același sunet în română, mapează la 'a' pentru matching (ex: cîmp/câmp -> camp)
    # Diacritice din alte limbi (pentru compatibilitate)
    'á': 'a', 'à': 'a', 'ä': 'a',
    'é': 'e', 'è': 'e', 'ê': 'e', 'ë': 'e',
    'í': 'i', 'ì': 'i', 'ï': 'i',
    'ó': 'o', 'ò': 'o', 'ô': 'o', 'ö': 'o',
    'ú': 'u', 'ù': 'u', 'û': 'u', 'ü': 'u',
}

# Reverse mapping: pentru fiecare literă de bază, toate variantele posibile
# Folosit pentru a genera toate variantele unui cuvânt de căutare
LETTER_VARIANTS = {
    'a': ['a', 'â', 'î', 'á', 'à', 'ä'],
    'e': ['e', 'é', 'è', 'ê', 'ë'],
    'i': ['i', 'í', 'ì', 'ï'],
    'o': ['o', 'ó', 'ò', 'ô', 'ö'],
    'u': ['u', 'ú', 'ù', 'û', 'ü'],
    's': ['s', 'ș'],
    't': ['t', 'ț'],
}


def normalize_similar_letters(s: str) -> str:
    """
    Normalizează litere similare la forma de bază.
    Ex: ochișorii -> ochisorii, cîmp -> camp, câmp -> camp
    Această funcție normalizează atât query-ul cât și textul OCR,
    astfel încât toate variantele să se potrivească.
    """
    result = []
    for char in s:
        # Folosim mapping-ul sau păstrăm caracterul original
        result.append(SIMILAR_LETTERS_MAP.get(char, char))
    return ''.join(result)


def get_letter_variants(char: str) -> list:
    """
    Returnează toate variantele posibile pentru o literă.
    Ex: 'a' -> ['a', 'â', 'î', 'á', 'à', 'ä']
        's' -> ['s', 'ș']
    """
    return LETTER_VARIANTS.get(char.lower(), [char.lower()])


def expand_word_variants(word: str) -> list:
    """
    Generează toate variantele posibile ale unui cuvânt prin înlocuirea
    literelor cu variantele lor (diacritice).
    Ex: 'camp' -> ['camp', 'câmp', 'cîmp', 'cámp', etc.]
    
    Notă: Poate genera multe variante, deci folosim normalizarea în loc
    de expansiune explicită pentru performanță.
    """
    # Pentru performanță, folosim normalizarea în loc de generarea tuturor variantelor
    # Această funcție este păstrată pentru referință, dar nu este folosită în căutare
    if not word:
        return ['']
    
    variants = ['']
    for char in word:
        char_variants = get_letter_variants(char)
        new_variants = []
        for variant in variants:
            for char_var in char_variants:
                new_variants.append(variant + char_var)
        variants = new_variants
    
    return variants


def normalize_text(s: str) -> str:
    """
    Normalizare simplă: litere mici, fără punctuație, spații compacte.
    Aplică și normalizarea literelor similare pentru matching flexibil.
    """
    s = s.lower()
    s = re.sub(r"[^\w\săâîșțáéíóúàèìòùâêîôûäëïöü-]", " ", s, flags=re.UNICODE)
    s = re.sub(r"\s+", " ", s).strip()
    # Aplică normalizarea literelor similare
    s = normalize_similar_letters(s)
    return s


def extract_page_number_from_image(img_name: str) -> str:
    """
    Extract the first full number from image name, removing leading zeros.
    Examples:
        "PPR-102.jpg" -> "102"
        "pag235.jpg" -> "235"
        "IMG_0520.jpg" -> "520"
        "pag172-173.jpg" -> "172"
    """
    if not img_name:
        return ""
    
    # Find first sequence of digits
    match = re.search(r'\d+', img_name)
    if match:
        number = match.group(0)
        # Remove leading zeros
        return number.lstrip('0') or '0'  # Keep '0' if all digits are zeros
    return ""


# ---------- OCR.TXT: PAGINI ----------

def load_ocr_pages(ocr_root: Path):
    """
    Din toate *.txt din subfolderele lui ocr_root extrage paginile OCR.

    Returnează listă de dict:
    {
        "folder":   numele subfolderului,
        "txt_file": numele fișierului txt,
        "page_num": numărul paginii (string),
        "page_img": numele imaginii (pagXXX.jpg),
        "text":     conținutul paginii,
        "lines":    lista de linii (pentru context)
    }
    """
    pages = []

    for folder in sorted(p for p in ocr_root.iterdir() if p.is_dir()):
        for txt in sorted(folder.glob("*.txt")):
            with txt.open("r", encoding="utf-8") as f:
                current_page_num = None
                current_page_img = None
                buffer = []
                lines_buffer = []

                def flush_page():
                    nonlocal buffer, current_page_num, current_page_img, lines_buffer
                    if current_page_num is None or current_page_img is None:
                        return
                    content = "".join(buffer).strip()
                    norm = normalize_text(content)
                    word_count = len(norm.split())
                    # aruncăm paginile cu prea puține cuvinte (zgomot OCR)
                    if word_count < MIN_PAGE_WORDS:
                        return
                    # Clean lines: remove empty lines and header markers
                    clean_lines = [line.strip() for line in lines_buffer if line.strip() and not PAGE_HEADER_RE.match(line.strip())]
                    # Extract page number from image name (first full number)
                    page_num_from_img = extract_page_number_from_image(current_page_img)
                    pages.append({
                        "folder": folder.name,
                        "txt_file": txt.name,
                        "page_num": page_num_from_img if page_num_from_img else current_page_num,
                        "page_img": current_page_img,
                        "text": content,
                        "lines": clean_lines,
                    })

                for line in f:
                    m = PAGE_HEADER_RE.match(line)
                    if m:
                        flush_page()
                        current_page_num = m.group(1)
                        current_page_img = m.group(2)
                        buffer = []
                        lines_buffer = []
                    else:
                        buffer.append(line)
                        lines_buffer.append(line)
                flush_page()

    print(f"Loaded {len(pages)} usable OCR pages from {ocr_root}")
    return pages


# ---------- SEARCH ----------

def find_match_context(query_text: str, page_lines: list, before: int = 2, after: int = 2):
    """
    Găsește poziția match-ului în linii și returnează contextul
    (before linii înainte, linia cu match, after linii după).
    """
    norm_query = normalize_text(query_text)
    query_words = norm_query.split()
    
    # Căutăm prima linie care conține majoritatea cuvintelor din query
    best_line_idx = None
    best_match_score = 0
    
    for i, line in enumerate(page_lines):
        norm_line = normalize_text(line)
        if not norm_line:
            continue
        
        # Numărăm câte cuvinte din query apar în linie
        words_found = sum(1 for word in query_words if word in norm_line)
        if words_found > 0:
            score = fuzz.partial_ratio(norm_query, norm_line)
            if score > best_match_score:
                best_match_score = score
                best_line_idx = i
    
    if best_line_idx is None:
        # Fallback: return first few lines if no match found
        return page_lines[:before + after + 1] if page_lines else []
    
    # Extragem contextul
    start_idx = max(0, best_line_idx - before)
    end_idx = min(len(page_lines), best_line_idx + after + 1)
    
    return page_lines[start_idx:end_idx]


def search_text_in_pages(query_text: str, pages, threshold: int = DEFAULT_THRESHOLD):
    """
    Caută query_text în toate paginile și returnează toate match-urile
    cu score >= threshold.
    
    IMPORTANT: Căutarea găsește automat TOATE variantele literelor:
    - "camp" va găsi: "camp", "cîmp", "câmp"
    - "ochisorii" va găsi: "ochisorii", "ochișorii"
    - "cîmp" va găsi: "camp", "cîmp", "câmp"
    - "câmp" va găsi: "camp", "cîmp", "câmp"
    
    Acest lucru se realizează prin normalizarea atât a query-ului cât și a textului OCR,
    astfel încât toate variantele diacritice se mapează la aceeași formă de bază.
    
    Returnează listă de dict:
    {
        "folder":   numele subfolderului,
        "image":    numele imaginii,
        "score":    score-ul match-ului (0-100),
        "page_num": numărul paginii,
        "snippet":  lista de linii cu context (2 înainte, 2 după)
    }
    """
    # Normalizează query-ul pentru a găsi toate variantele
    # Ex: "camp" -> "camp", "cîmp" -> "camp", "câmp" -> "camp"
    norm_query = normalize_text(query_text)
    query_words = norm_query.split()
    
    # Validare: query trebuie să aibă între 2 și 12 cuvinte
    if len(query_words) < 2:
        raise ValueError("Query must have at least 2 words")
    if len(query_words) > 12:
        raise ValueError("Query must have at most 12 words")
    
    matches = []
    
    for page in pages:
        # Normalizează textul OCR pentru a se potrivi cu toate variantele din query
        # Ex: "cîmp" -> "camp", "câmp" -> "camp", "camp" -> "camp"
        norm_page = normalize_text(page["text"])
        
        if not norm_page:
            continue
        
        # Folosim partial_ratio pentru a găsi query-ul în text
        # (funcționează bine pentru substring matching)
        score = fuzz.partial_ratio(norm_query, norm_page)
        
        # De asemenea, verificăm dacă toate cuvintele din query apar în pagină
        # pentru a evita false positives
        words_found = sum(1 for word in query_words if word in norm_page)
        word_coverage = (words_found / len(query_words)) * 100 if query_words else 0
        
        # Combinăm score-urile: partial_ratio și word coverage
        # Dacă toate cuvintele sunt prezente, boostăm score-ul
        if words_found == len(query_words):
            # Toate cuvintele sunt prezente - boost score
            final_score = max(score, word_coverage * 0.9)
        elif words_found >= len(query_words) * 0.8:
            # Cel puțin 80% din cuvinte sunt prezente
            final_score = (score + word_coverage) / 2
        else:
            # Prea puține cuvinte - folosim doar partial_ratio
            final_score = score * (words_found / len(query_words))
        
        if final_score >= threshold:
            # Găsim contextul pentru snippet
            snippet = find_match_context(query_text, page.get("lines", []))
            
            matches.append({
                "folder": page["folder"],
                "image": page["page_img"],
                "score": round(final_score, 1),
                "page_num": page["page_num"],
                "snippet": snippet
            })
    
    # Sortăm după score descrescător
    matches.sort(key=lambda x: x["score"], reverse=True)
    
    return matches


# ---------- MAIN ----------

def main():
    if len(sys.argv) < 3:
        print("Usage: python search_text.py \"text query\" OCR_ROOT [threshold] [output_dir]")
        print("  text query = 2-12 word text to search for")
        print("  OCR_ROOT   = folder with OCR .txt files")
        print("  threshold  = minimum score (0-100), default 70")
        print("  output_dir = directory to save default.docx (optional)")
        sys.exit(1)

    query_text = sys.argv[1]
    ocr_root = Path(sys.argv[2])
    threshold = int(sys.argv[3]) if len(sys.argv) > 3 else DEFAULT_THRESHOLD
    output_dir = Path(sys.argv[4]) if len(sys.argv) > 4 else Path.cwd()

    if not ocr_root.is_dir():
        print(f"OCR root folder not found: {ocr_root}")
        sys.exit(1)

    if threshold < 0 or threshold > 100:
        print(f"Threshold must be between 0 and 100, got: {threshold}")
        sys.exit(1)

    print(f"Loading OCR pages from {ocr_root} ...")
    pages = load_ocr_pages(ocr_root)

    print(f"\nSearching for: \"{query_text}\"")
    print(f"Threshold: {threshold}")
    print("-" * 60)

    try:
        matches = search_text_in_pages(query_text, pages, threshold)
        
        # Write results to text file
        output_file = output_dir / "search-result.txt"
        with output_file.open("w", encoding="utf-8") as f:
            if not matches:
                f.write(f"No matches found above threshold {threshold}\n")
                print(f"No matches found above threshold {threshold}")
            else:
                print(f"Found {len(matches)} match(es)")
                print(f"Results saved to: {output_file}\n")
                
                for i, match in enumerate(matches, 1):
                    # Write verses
                    for snippet_line in match['snippet']:
                        if snippet_line.strip():  # Only write non-empty lines
                            f.write(snippet_line + "\n")
                    
                    # Write source in parentheses
                    f.write(f"({match['folder']}, p. {match['page_num']})\n")
                    f.write("\n")
                    
                    # Also print to console
                    print(f"{i}. Score: {match['score']:.1f} | {match['folder']} | {match['image']} | Page {match['page_num']}")
        
        # Write results to Word document
        # Use output_dir if provided, otherwise use current working directory
        docx_file = output_dir / "default.docx"
        if docx_file.exists():
            doc = Document(str(docx_file))
            # Clear existing content (keep styles) by removing all paragraphs
            for para in list(doc.paragraphs):
                p = para._element
                p.getparent().remove(p)
        else:
            doc = Document()
        
        # Add title with search term in italics and quotes
        title_para = doc.add_paragraph()
        title_run = title_para.add_run('Search results for "')
        search_run = title_para.add_run(query_text)
        search_run.italic = True
        title_para.add_run('"')
        
        # Add spacing after title
        doc.add_paragraph()
        
        if not matches:
            para = doc.add_paragraph(f"No matches found above threshold {threshold}")
        else:
            for i, match in enumerate(matches, 1):
                # Add verses with style "2-Versuri-centru"
                for snippet_line in match['snippet']:
                    if snippet_line.strip():  # Only add non-empty lines
                        para = doc.add_paragraph(snippet_line.strip())
                        try:
                            para.style = "2-Versuri-centru"
                        except KeyError:
                            print(f"  Warning: Style '2-Versuri-centru' not found, using default")
                
                # Add source with style "4-Sursa text"
                source_text = f"({match['folder']}, p. {match['page_num']})"
                para = doc.add_paragraph(source_text)
                try:
                    para.style = "4-Sursa text"
                except KeyError:
                    print(f"  Warning: Style '4-Sursa text' not found, using default")
                
                # Add empty paragraph for spacing
                doc.add_paragraph()
        
        doc.save(str(docx_file))
        print(f"Results also saved to: {docx_file}")
        
    except ValueError as e:
        print(f"Error: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()

