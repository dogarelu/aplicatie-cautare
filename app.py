import tkinter as tk
from tkinter import messagebox, filedialog, ttk
import shutil
import sys
import re
import subprocess
import platform
import requests
import time
import os
import multiprocessing
import json
from typing import List, Set, Dict
from pathlib import Path
from rapidfuzz import fuzz
from rapidfuzz.distance import Levenshtein
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from concurrent.futures import ThreadPoolExecutor, ProcessPoolExecutor, as_completed
import threading
try:
    from PIL import Image, ImageTk
except ImportError:
    Image = None
    ImageTk = None
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None


# Global variable to store selected root folder and selected paths
selected_root_folder = None
selected_paths = []  # List of selected paths (Biblioteca, Rafturi, or Volume paths)
selected_volumes = []  # List of selected "volumes" (final folders with ocr.txt)
status_var = None

# ---------- OCR Constants ----------
# Header de pagină în fișierele OCR (*.txt)
# ex: === PAGE 151 (pag151.jpg) ===
PAGE_HEADER_RE = re.compile(r"^=== PAGE (.+?) \((.+?)\) ===")

# Pagini cu foarte puține cuvinte sunt ignorate (zgomot)
MIN_PAGE_WORDS = 6

# Default threshold pentru matching
DEFAULT_THRESHOLD = 80

# ---------- Text Normalization ----------
# Mapping pentru litere similare (pentru căutare OCR)
# Mapează diacritice la forma de bază pentru matching flexibil
SIMILAR_LETTERS_MAP = {
    # Diacritice românești
    'ș': 's',
    'ț': 't',
    'â': 'a',
    'î': 'a',  # î și â sunt același sunet în română, mapează la 'a' pentru matching (ex: cîmp/câmp -> camp)
    # Diacritice din alte limbi (pentru compatibilitate)
    'á': 'a', 'à': 'a', 'ä': 'a',
    'é': 'e', 'è': 'e', 'ê': 'e', 'ë': 'e',
    'í': 'i', 'ì': 'i', 'ï': 'i',
    'ó': 'o', 'ò': 'o', 'ô': 'o', 'ö': 'o',
    'ú': 'u', 'ù': 'u', 'û': 'u', 'ü': 'u',
}

# Reverse mapping: pentru fiecare literă de bază, toate variantele posibile
LETTER_VARIANTS = {
    'a': ['a', 'â', 'î', 'á', 'à', 'ä'],
    'e': ['e', 'é', 'è', 'ê', 'ë'],
    'i': ['i', 'í', 'ì', 'ï'],
    'o': ['o', 'ó', 'ò', 'ô', 'ö'],
    'u': ['u', 'ú', 'ù', 'û', 'ü'],
    's': ['s', 'ș'],
    't': ['t', 'ț'],
}


def get_base_path():
    """
    Get the base directory path that works with both script and PyInstaller executable.
    Returns the directory where the executable/script is located.
    For PyInstaller --onefile mode, returns the temp extraction directory (for finding bundled files).
    For PyInstaller --onedir mode, returns the directory containing the executable.
    """
    if getattr(sys, 'frozen', False):
        # Running as compiled executable (PyInstaller)
        # Check if we're in onefile mode (temp extraction directory exists)
        if hasattr(sys, '_MEIPASS'):
            # --onefile mode: bundled files are in temp directory
            return Path(sys._MEIPASS)
        else:
            # --onedir mode: files are next to executable
            return Path(sys.executable).parent
    else:
        # Running as script
        return Path(__file__).parent


def get_output_path():
    """
    Get the directory where output files should be saved.
    Always returns a user-accessible location (executable directory, not temp).
    """
    if getattr(sys, 'frozen', False):
        # Always save to executable directory (user-accessible)
        return Path(sys.executable).parent
    else:
        # Running as script
        return Path(__file__).parent


def get_app_config():
    """Get app configuration from file."""
    config_file = get_output_path() / ".app_config.json"
    if config_file.exists():
        try:
            with open(config_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        except:
            pass
    return {}

def save_app_config(config: dict):
    """Save app configuration to file."""
    config_file = get_output_path() / ".app_config.json"
    try:
        with open(config_file, 'w', encoding='utf-8') as f:
            json.dump(config, f, indent=2)
    except:
        pass

def get_last_library_dir():
    """Get the last library directory from config file."""
    config = get_app_config()
    last_dir = config.get('last_library_dir')
    if last_dir and Path(last_dir).exists():
        return last_dir
    return None

def save_last_library_dir(directory: str):
    """Save the last library directory to config file."""
    config = get_app_config()
    config['last_library_dir'] = directory
    save_app_config(config)

def get_last_export_dir():
    """Get the last export directory from config file."""
    config = get_app_config()
    last_dir = config.get('last_export_dir')
    if last_dir and Path(last_dir).exists():
        return last_dir
    return None

def save_last_export_dir(directory: str):
    """Save the last export directory to config file."""
    config = get_app_config()
    config['last_export_dir'] = directory
    save_app_config(config)


# ---------- Text Normalization Functions ----------

def normalize_similar_letters(s: str) -> str:
    """
    Normalizează litere similare la forma de bază.
    Ex: ochișorii -> ochisorii, cîmp -> camp, câmp -> camp
    """
    result = []
    for char in s:
        result.append(SIMILAR_LETTERS_MAP.get(char, char))
    return ''.join(result)


def normalize_text(s: str) -> str:
    """
    Normalizare simplă: litere mici, fără punctuație, spații compacte.
    Aplică și normalizarea literelor similare pentru matching flexibil.
    """
    s = s.lower()
    s = re.sub(r"[^\w\săâîșțáéíóúàèìòùâêîôûäëïöü-]", " ", s, flags=re.UNICODE)
    s = re.sub(r"\s+", " ", s).strip()
    s = normalize_similar_letters(s)
    return s


def words_match_simple(query_word: str, text_word: str) -> bool:
    """
    Simple word matching without hyphen handling (used internally).
    """
    if not query_word or not text_word:
        return False
    
    # Exact match always works
    if query_word == text_word:
        return True
    
    # For words 3 characters or less, require exact match
    if len(query_word) <= 3 or len(text_word) <= 3:
        return False
    
    # Calculate Levenshtein distance using rapidfuzz
    lev_distance = Levenshtein.distance(query_word, text_word)
    
    # Determine allowed distance based on word length
    # Use the longer of the two words to determine the threshold
    max_length = max(len(query_word), len(text_word))
    
    # For words 4-7 characters, allow 1 character difference
    if 4 <= max_length <= 7:
        return lev_distance <= 1
    
    # For words >= 8 characters, allow 2 character differences
    if max_length >= 8:
        return lev_distance <= 2
    
    # Fallback (shouldn't reach here, but handle it)
    return lev_distance <= 1


def words_match(query_word: str, text_word: str) -> bool:
    """
    Check if two words match according to fuzzy matching rules:
    - If word length <= 3: must be exact match
    - If word length 4-7: allow 1 character difference (Levenshtein distance <= 1)
    - If word length >= 8: allow 2 character differences (Levenshtein distance <= 2)
    - Similar letters/variants don't count as different (already normalized)
    - Handles hyphenated words: matches parts before/after hyphens
    
    Args:
        query_word: The search word (normalized)
        text_word: The word from text to match against (normalized)
    
    Returns:
        True if words match according to the rules
    """
    # Both words should already be normalized (similar letters handled)
    if not query_word or not text_word:
        return False
    
    # Exact match always works
    if query_word == text_word:
        return True
    
    # Handle hyphenated words: check if query matches any part of hyphenated text word
    # e.g., "calu" should match "calu-m" (part before hyphen)
    if '-' in text_word:
        text_parts = text_word.split('-')
        for part in text_parts:
            if part and words_match_simple(query_word, part):
                return True
    
    # Handle hyphenated query: check if any part of hyphenated query matches text word
    # e.g., "intr-o" should match "intro"
    if '-' in query_word:
        query_parts = query_word.split('-')
        for part in query_parts:
            if part and words_match_simple(part, text_word):
                return True
    
    # Also check if removing hyphens makes them match
    # e.g., "intro" should match "intr-o"
    text_word_no_hyphen = text_word.replace('-', '')
    query_word_no_hyphen = query_word.replace('-', '')
    if text_word_no_hyphen and query_word_no_hyphen:
        if words_match_simple(query_word_no_hyphen, text_word_no_hyphen):
            return True
    
    # Standard matching for non-hyphenated words
    return words_match_simple(query_word, text_word)


def get_inflected_forms_from_dex(word: str, cache: Dict[str, Set[str]] = None) -> Set[str]:
    """
    Obține formele flexionare ale unui cuvânt folosind API-ul DEX online.
    
    Args:
        word: Cuvântul pentru care se caută forme flexionare
        cache: Dicționar pentru cache (opțional, pentru a evita cereri duplicate)
    
    Returns:
        Set de forme flexionare (include și cuvântul original)
    """
    if cache is None:
        cache = {}
    
    word_lower = word.lower().strip()
    
    # Verifică cache-ul
    if word_lower in cache:
        return cache[word_lower]
    
    forms = {word_lower}  # Include forma originală
    
    try:
        # Endpoint API DEX online pentru lexem
        url = f"https://dexonline.ro/api/lexem/{word_lower}"
        
        response = requests.get(url, timeout=5)
        
        if response.status_code == 200:
            data = response.json()
            
            # Extrage formele flexionare din răspuns
            # Structura poate varia - ajustează în funcție de răspunsul real
            if isinstance(data, dict):
                # Caută câmpuri care conțin forme flexionare
                if 'flexiuni' in data:
                    flexiuni = data['flexiuni']
                    if isinstance(flexiuni, list):
                        for form in flexiuni:
                            if isinstance(form, str):
                                forms.add(form.lower())
                            elif isinstance(form, dict) and 'form' in form:
                                forms.add(form['form'].lower())
                
                # Poate fi și în alt format
                if 'forms' in data:
                    forms_data = data['forms']
                    if isinstance(forms_data, list):
                        for form in forms_data:
                            if isinstance(form, str):
                                forms.add(form.lower())
                            elif isinstance(form, dict) and 'form' in form:
                                forms.add(form['form'].lower())
                
                # Sau în definiții
                if 'definitions' in data:
                    for def_item in data['definitions']:
                        if isinstance(def_item, dict) and 'flexiuni' in def_item:
                            flexiuni = def_item['flexiuni']
                            if isinstance(flexiuni, list):
                                for form in flexiuni:
                                    if isinstance(form, str):
                                        forms.add(form.lower())
                
                # Caută și în structura de lexem
                if 'lexem' in data:
                    lexem_data = data['lexem']
                    if isinstance(lexem_data, dict) and 'flexiuni' in lexem_data:
                        flexiuni = lexem_data['flexiuni']
                        if isinstance(flexiuni, list):
                            for form in flexiuni:
                                if isinstance(form, str):
                                    forms.add(form.lower())
            
        # Rate limiting - așteaptă puțin între cereri
        time.sleep(0.1)
        
    except requests.exceptions.RequestException as e:
        # Dacă API-ul nu funcționează, returnează doar forma originală
        pass  # Fail silently - use original word only
    except Exception as e:
        pass  # Fail silently - use original word only
    
    # Salvează în cache
    cache[word_lower] = forms
    
    return forms


def expand_search_terms_with_inflections(query_text: str, cache: Dict[str, Set[str]] = None) -> List[str]:
    """
    Extinde termenii de căutare cu formele lor flexionare.
    
    Args:
        query_text: Textul de căutare original
        cache: Cache pentru forme flexionare (opțional)
    
    Returns:
        Listă de cuvinte extinse (original + forme flexionare)
    """
    words = query_text.split()
    expanded_words = []
    
    for word in words:
        # Obține forme flexionare
        inflected_forms = get_inflected_forms_from_dex(word, cache)
        expanded_words.extend(inflected_forms)
    
    # Returnează lista unică, păstrând ordinea
    seen = set()
    result = []
    for word in expanded_words:
        if word not in seen:
            seen.add(word)
            result.append(word)
    
    return result


def extract_page_number_from_image(img_name: str) -> str:
    """
    Extract the first full number from image name, removing leading zeros.
    Examples:
        "PPR-102.jpg" -> "102"
        "pag235.jpg" -> "235"
        "IMG_0520.jpg" -> "520"
        "pag172-173.jpg" -> "172"
    """
    if not img_name:
        return ""
    match = re.search(r'\d+', img_name)
    if match:
        number = match.group(0)
        return number.lstrip('0') or '0'
    return ""


# ---------- Table of Contents Functions ----------


def _sanitize_title(raw_title: str) -> str:
    """Normalize titles from cuprins: drop numbering, parenthetical notes, and keep only letters/digits/spaces/-/'."""
    if not raw_title:
        return ""
    title = raw_title.strip()
    # Remove leading numbering like "12." or "12)" or "12 -"
    title = re.sub(r'^\s*\d+[.)]?\s*', '', title)
    # Remove any parenthetical content e.g., "(cdnsjad)"
    title = re.sub(r'\([^)]*\)', '', title)
    # Keep only unicode letters/digits, spaces, hyphen, apostrophe (preserve diacritics)
    title = re.sub(r"[^\w\s\-']", " ", title, flags=re.UNICODE)
    title = title.replace("_", " ")
    # Collapse whitespace
    title = " ".join(title.split())
    return title.strip()

def load_cuprins(volume_path: Path) -> Dict[int, str]:
    """
    Load and parse cuprins.txt file from a volume folder.
    Returns a dictionary mapping page numbers to titles.
    Format: "TITLE -- PAGE_NUMBER"
    """
    cuprins_file = volume_path / "cuprins.txt"
    if not cuprins_file.exists():
        return {}
    
    toc = {}  # page_num -> title
    try:
        with cuprins_file.open("r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                # Parse format: "TITLE -- PAGE_NUMBER"
                # Handle different separators: --, -, or just spaces
                if " -- " in line:
                    parts = line.split(" -- ", 1)
                elif " - " in line:
                    parts = line.split(" - ", 1)
                else:
                    # Try to find last number in the line
                    match = re.search(r'(\d+)\s*$', line)
                    if match:
                        page_num_str = match.group(1)
                        title = _sanitize_title(line[:match.start()])
                        try:
                            page_num = int(page_num_str)
                            if title:
                                toc[page_num] = title
                        except ValueError:
                            pass
                    continue
                
                if len(parts) == 2:
                    title = _sanitize_title(parts[0])
                    page_num_str = parts[1].strip()
                    try:
                        page_num = int(page_num_str)
                        if title:
                            toc[page_num] = title
                    except ValueError:
                        pass
    except Exception as e:
        print(f"[{time.strftime('%H:%M:%S')}] Error loading cuprins.txt from {volume_path}: {e}")
    
    return toc


def find_title_for_page(volume_path: Path, page_num: int, page_text: str = None, match_line_idx: int = None, page_lines: list = None) -> str:
    """
    Find the title for a given page number.
    If page_num is exactly on a boundary, check OCR text to see if title appears before or after the match.
    
    Args:
        volume_path: Path to the volume folder (contains cuprins.txt)
        page_num: Page number where the match was found
        page_text: Full text of the page (for boundary checking)
        match_line_idx: Line index where the match occurs (for boundary checking)
        page_lines: List of page lines (for boundary checking)
    
    Returns:
        Title string or "qqq" if not found
    """
    toc = get_cuprins_cached(volume_path)
    if not toc:
        return "qqq"
    
    # Convert page_num to int if it's a string
    try:
        if isinstance(page_num, str):
            page_num = page_num.strip()
            if not page_num or page_num == "0":
                return "qqq"
            page_num_int = int(page_num)
        else:
            page_num_int = int(page_num)
    except (ValueError, TypeError):
        return "qqq"
    
    # Find the title that should contain this page
    # Sort by page number
    sorted_pages = sorted(toc.keys())
    
    # Check if there's a title that starts exactly at this page
    title_at_page = toc.get(page_num_int)
    
    # Find the last title that starts before this page (previous title)
    previous_title = ""
    previous_title_start_page = None
    
    for start_page in sorted_pages:
        if start_page < page_num_int:
            previous_title = toc[start_page]
            previous_title_start_page = start_page
        else:
            break
    
    # If there's a title that starts at this page, check OCR text to determine which title to use
    if title_at_page and page_text and match_line_idx is not None and page_lines:
        norm_page_text = normalize_text(page_text)
        norm_match_text = normalize_text(" ".join(page_lines[:match_line_idx + 1]))
        match_position = len(norm_match_text)
        
        # Check for current title in the page text
        norm_current_title = normalize_text(title_at_page)
        current_title_position = None
        if norm_current_title in norm_page_text:
            current_title_position = norm_page_text.find(norm_current_title)
        else:
            current_title_words = norm_current_title.split()
            if current_title_words:
                first_word = current_title_words[0]
                if first_word in norm_page_text:
                    current_title_position = norm_page_text.find(first_word)
        
        # Decide based only on current vs previous title
        if current_title_position is not None:
            # Current title before or at match => keep current title
            if current_title_position <= match_position:
                return title_at_page
            # Current title appears after match => use previous title if available
            if previous_title:
                return previous_title
            return title_at_page
        
        # If current title not found in OCR text, default to current title
        return title_at_page
    
    # If no title starts at this page, use the previous title
    if previous_title:
        return previous_title
    
    # If no title found, return "qqq"
    return "qqq"
    
    # Return title if found, otherwise return "qqq" as fallback
    return title if title else "qqq"


# Cache for cuprins files to avoid reloading
_cuprins_cache = {}


def get_cuprins_cached(volume_path: Path) -> Dict[int, str]:
    """Get cuprins with caching."""
    volume_path_str = str(volume_path)
    if volume_path_str not in _cuprins_cache:
        _cuprins_cache[volume_path_str] = load_cuprins(volume_path)
    return _cuprins_cache[volume_path_str]


# ---------- Index RTF Functions ----------

def extract_rtf_text(rtf_content: str) -> str:
    """
    Extract plain text from RTF content by removing RTF control codes.
    Handles RTF escape sequences for special characters.
    """
    # First, handle RTF escape sequences for special characters
    # \'XX where XX is hex code for character
    def replace_rtf_escape(match):
        hex_code = match.group(1)
        try:
            char_code = int(hex_code, 16)
            return chr(char_code)
        except:
            return match.group(0)
    
    # Replace \'XX escape sequences
    text = re.sub(r"\\'([0-9a-fA-F]{2})", replace_rtf_escape, rtf_content)
    
    # Remove RTF control words (but preserve text)
    # Remove \ commands that are control words
    text = re.sub(r'\\[a-z]+\d*\s*', ' ', text)
    # Remove RTF groups that are just formatting
    # Remove braces but be careful - we'll do it in a way that preserves text
    # Remove standalone braces
    text = re.sub(r'[{}]', ' ', text)
    # Remove special RTF sequences
    text = re.sub(r'\\[\\\-\{\}]', '', text)
    # Remove remaining RTF control sequences
    text = re.sub(r'\\[^a-zA-Z\s]', '', text)
    # Clean up multiple spaces and newlines
    text = re.sub(r'\s+', ' ', text)
    # Replace \par with newline
    text = text.replace('\\par', '\n')
    # Clean up
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


def load_index_rtf(index_file: Path) -> Dict[str, List[str]]:
    """
    Load and parse an index file (BF-index.txt or CEE-index.txt).
    Returns a dictionary mapping normalized titles to lists of codes.
    Format: "Title, CODE1; CODE2; CODE3" -> {"title": ["CODE1", "CODE2", "CODE3"]}
    Handles formats like "97.III", "130.I subtip", "131.II. 2" - extracts base number.
    """
    if not index_file.exists():
        return {}
    
    load_start = time.perf_counter()
    index_dict = {}
    try:
        with index_file.open("r", encoding="utf-8", errors="ignore") as f:
            for line in f:
                line = line.strip()
                if not line or len(line) < 3:
                    continue
                
                # Skip header lines
                if line.startswith("Amzulescu") or "INDICE" in line or "ALFABETIC" in line:
                    continue
                
                # Split by comma to separate title from codes
                # Format: "Title, CODE1; CODE2; CODE3" or "Title, CODE1"
                if ',' not in line:
                    continue
                
                parts = line.split(',', 1)
                if len(parts) != 2:
                    continue
                
                title_clean = parts[0].strip()
                codes_part = parts[1].strip()
                
                # Only process if we have a reasonable title (at least 3 characters)
                if len(title_clean) < 3:
                    continue
                
                # Extract all codes from the codes part
                # Codes can be separated by semicolons: "97.III; 130.I subtip"
                # Extract full code format (including dots, Roman numerals, and "subtip" if present)
                codes = []
                # Split by semicolon to get individual codes
                code_parts = re.split(r'[;]', codes_part)
                for code_part in code_parts:
                    code_part = code_part.strip()
                    # Extract the full code format (digits, dots, letters for Roman numerals, and "subtip" if present)
                    # Examples: "5.I" -> "5.I", "97.III" -> "97.III", "130.I subtip" -> "130.I subtip", "131.II. 2" -> "131.II"
                    # Match: digits, optional dots with letters, optional space followed by "subtip"
                    match = re.search(r'^(\d+(?:\.[A-Za-z]+)*(?:\s+subtip)?)', code_part, re.IGNORECASE)
                    if match:
                        code_num = match.group(1).strip()
                        codes.append(code_num)
                
                if codes:
                    # Normalize title for matching (normalize_text already lowercases)
                    norm_title = normalize_text(title_clean)
                    if norm_title:
                        # Store all codes for this title
                        if norm_title not in index_dict:
                            index_dict[norm_title] = []
                        # Add codes that aren't already in the list
                        for code in codes:
                            if code not in index_dict[norm_title]:
                                index_dict[norm_title].append(code)
        
        load_end = time.perf_counter()
        load_time = load_end - load_start
        print(f"[{time.strftime('%H:%M:%S')}] Loaded index file '{index_file.name}': {len(index_dict)} entries in {load_time:.2f}s")
    except Exception as e:
        print(f"[{time.strftime('%H:%M:%S')}] Error loading index file {index_file}: {e}")
    
    return index_dict


def load_coduri_doc() -> Dict[str, str]:
    """
    Load Coduri-BF-CEE.txt and return a dictionary mapping codes to full lines.
    Format: {"BF 5": "BF 5 - Voinicul rănit  • Tip I - [Maica bătrână] 7", "CEE 16": "CEE 16 - Voinicul adormit 221"}
    Returns the first line found for each code (in case there are multiple types).
    """
    app_root = get_output_path()
    coduri_file = app_root / "Coduri-BF-CEE.txt"
    
    # Also try .doc extension in case it's actually a .doc file
    if not coduri_file.exists():
        coduri_file = app_root / "Coduri-BF-CEE.doc"
    
    # Also try .docx extension
    if not coduri_file.exists():
        coduri_file = app_root / "Coduri-BF-CEE.docx"
    
    if not coduri_file.exists():
        print(f"[{time.strftime('%H:%M:%S')}] Coduri-BF-CEE file not found in {app_root}")
        return {}
    
    coduri_dict = {}
    try:
        # Try reading as .docx first (python-docx only supports .docx)
        if coduri_file.suffix == ".docx":
            try:
                doc = Document(str(coduri_file))
                for para in doc.paragraphs:
                    line = para.text.strip()
                    if line:
                        # Look for pattern: "BF 74 - TITLE" or "CEE 16 - TITLE"
                        # Also handle codes with suffixes like "BF 5.I" (though they may not exist in this file)
                        match = re.match(r'^(BF|CEE)\s+(\d+(?:\.[A-Za-z]+)?)\s*-\s*(.+)$', line, re.IGNORECASE)
                        if match:
                            code_type = match.group(1).upper()
                            code_num = match.group(2)  # This can be "5" or "5.I" if it exists
                            code = f"{code_type} {code_num}"
                            # Only store first occurrence of each code
                            if code not in coduri_dict:
                                coduri_dict[code] = line
            except Exception as e:
                print(f"[{time.strftime('%H:%M:%S')}] Error reading Coduri-BF-CEE as .docx: {e}")
        
        # Try reading as plain text (.txt or fallback)
        if coduri_file.suffix == ".txt" or coduri_file.suffix == ".doc" or not coduri_dict:
            try:
                with coduri_file.open("r", encoding="utf-8", errors="ignore") as f:
                    for line in f:
                        line = line.strip()
                        if line:
                            # Skip lines that don't start with BF or CEE
                            if not re.match(r'^(BF|CEE)\s+\d+', line, re.IGNORECASE):
                                continue
                            
                            # Look for pattern: "BF 74 - TITLE" or "CEE 16 - TITLE"
                            # Also handle lines with additional info like "BF 5 - Voinicul rănit  • Tip I - [Maica bătrână] 7"
                            # Also handle codes with suffixes like "BF 5.I" (though they may not exist in this file)
                            match = re.match(r'^(BF|CEE)\s+(\d+(?:\.[A-Za-z]+)?)\s*-\s*(.+)$', line, re.IGNORECASE)
                            if match:
                                code_type = match.group(1).upper()
                                code_num = match.group(2)  # This can be "5" or "5.I" if it exists
                                code = f"{code_type} {code_num}"
                                # Only store first occurrence of each code (first type)
                                if code not in coduri_dict:
                                    coduri_dict[code] = line
            except Exception as e2:
                print(f"[{time.strftime('%H:%M:%S')}] Error reading Coduri-BF-CEE as text: {e2}")
    except Exception as e:
        print(f"[{time.strftime('%H:%M:%S')}] Error loading Coduri-BF-CEE: {e}")
    
    print(f"[{time.strftime('%H:%M:%S')}] Loaded Coduri-BF-CEE: {len(coduri_dict)} codes")
    return coduri_dict


# Cache for coduri doc
_coduri_cache = None


def get_coduri_cached() -> Dict[str, str]:
    """Get coduri doc with caching."""
    global _coduri_cache
    if _coduri_cache is None:
        _coduri_cache = load_coduri_doc()
    return _coduri_cache


def find_code_for_title(volume_path: Path, title: str) -> str:
    """
    Find the code(s) for a title by searching in BF-index.txt and CEE-index.txt.
    If multiple matches found, returns all codes followed by "qqq".
    If single match found, looks up the full line from Coduri-BF-CEE.doc.
    
    Args:
        volume_path: Path to the volume folder (or parent folder containing index files)
        title: The title to search for
    
    Returns:
        Code string(s) like "BF 16" or "CEE 16", or "BF 16, CEE 20 qqq" for multiple,
        or full line from Coduri-BF-CEE.doc for single match
    """
    if not title or title == "qqq":
        return ""
    
    # Normalize title for matching (normalize_text already lowercases)
    norm_title = normalize_text(title)
    if not norm_title:
        return ""
    
    # Try to find index files - check app root folder first (where default.docx is),
    # then volume folder, then parent folders
    app_root = get_output_path()  # Where default.docx is located
    search_paths = [app_root, volume_path]
    # Also check parent folders up to 3 levels
    current = volume_path.parent
    for _ in range(3):
        if current and current != current.parent:
            search_paths.append(current)
            current = current.parent
        else:
            break
    
    search_start = time.perf_counter()
    
    # Determine max_diff based on title length (letters only)
    # Remove spaces and non-letters from normalized title to get length
    letters_only_title = re.sub(r'[^a-zăâîșțáéíóúàèìòùâêîôûäëïöü]', '', norm_title)
    # If title is 5 characters or less, allow no character differences (exact match only)
    # If title is 6-7 characters, allow 1 character difference
    # Otherwise, allow 2 character differences
    if len(letters_only_title) <= 5:
        max_diff = 0  # Exact match only
    elif len(letters_only_title) <= 7:
        max_diff = 1  # Allow 1 character difference
    else:
        max_diff = 2  # Allow 2 character differences
    
    # Helper function to check if two normalized strings match with maximum character differences
    # (ignoring spaces, non-letters, and diacritics)
    # The matched title can contain extra words as long as it contains the search title
    def matches_with_tolerance(str1: str, str2: str, max_diff: int = 2) -> tuple[bool, int, str, str]:
        """
        Check if two strings match, allowing up to max_diff character differences.
        Spaces, non-letters, and diacritics are normalized/ignored when calculating differences.
        The matched title (str2) must have at least as many words as the search title (str1).
        For max_diff=0 (5 letters or under), only exact matches are allowed (no character differences).
        Returns (matches, distance, letters_only_str1, letters_only_str2)
        """
        # Both should already be normalized and lowercase
        str1_lower = str1.lower()
        str2_lower = str2.lower()
        
        # Count words in each string (after normalization)
        words_1 = len([w for w in str1_lower.split() if w])
        words_2 = len([w for w in str2_lower.split() if w])
        
        # The matched title (str2) must have at least as many words as the search title (str1)
        if words_2 < words_1:
            return (False, 0, "", "")
        
        # Remove spaces and non-letters (keep only letters with diacritics)
        letters_only_1 = re.sub(r'[^a-zăâîșțáéíóúàèìòùâêîôûäëïöü]', '', str1_lower)
        letters_only_2 = re.sub(r'[^a-zăâîșțáéíóúàèìòùâêîôûäëïöü]', '', str2_lower)
        
        # Normalize diacritics (ă/â/î -> a, ș -> s, ț -> t, etc.) so diacritics don't count as differences
        letters_only_1_normalized = normalize_similar_letters(letters_only_1)
        letters_only_2_normalized = normalize_similar_letters(letters_only_2)
        
        # Exact match on letters only (after diacritic normalization)
        if letters_only_1_normalized == letters_only_2_normalized:
            return (True, 0, letters_only_1_normalized, letters_only_2_normalized)
        
        # If max_diff is 0 (5 letters or under), only allow exact matches - no substring matches with differences
        if max_diff == 0:
            return (False, 0, letters_only_1_normalized, letters_only_2_normalized)
        
        # For max_diff > 0, check if str2 contains str1 (allowing extra words/characters around it)
        # If str2 contains str1 as a substring (letters only, diacritics normalized), that's a valid match
        if letters_only_1_normalized in letters_only_2_normalized:
            # For substring matches, we allow the extra characters (they're just extra words)
            return (True, 0, letters_only_1_normalized, letters_only_2_normalized)
        
        # Calculate Levenshtein distance on letters only (after diacritic normalization)
        distance = Levenshtein.distance(letters_only_1_normalized, letters_only_2_normalized)
        
        # Allow up to max_diff character differences
        matches = distance <= max_diff
        return (matches, distance, letters_only_1_normalized, letters_only_2_normalized)
    
    # Collect all matching codes (not just the best one) from both BF and CEE
    all_matches = []  # List of (code, similarity, matched_title) tuples
    
    # Try BF-index.txt - search all paths
    bf_found = False
    for search_path in search_paths:
        bf_index = search_path / "BF-index.txt"
        if bf_index.exists():
            bf_dict = get_index_rtf_cached(bf_index)
            bf_found = True
            # First try exact match (already normalized)
            if norm_title in bf_dict:
                code_nums = bf_dict[norm_title]  # This is now a list of codes like ["5.I"]
                for code_num in code_nums:
                    code = f"BF {code_num}"
                    all_matches.append((code, 100.0, norm_title))
            # Always also try fuzzy matching to catch all possible matches
            # Try matching on all entries with dynamic character difference tolerance (ignoring spaces/non-letters)
            for dict_title, code_nums in bf_dict.items():
                # Skip if we already found exact match for this title
                if norm_title == dict_title:
                    continue
                # Both are already normalized, so we can compare directly
                matches, distance, letters_1, letters_2 = matches_with_tolerance(norm_title, dict_title, max_diff)
                
                if matches:
                    # Calculate similarity percentage
                    max_len = max(len(letters_1), len(letters_2))
                    similarity = (1.0 - distance / max_len) * 100.0 if max_len > 0 else 100.0
                    
                    # Add all codes for this matching title
                    for code_num in code_nums:
                        code = f"BF {code_num}"
                        all_matches.append((code, similarity, dict_title))
            break  # Only check first found file
    
    # Try CEE-index.txt - search all paths (always search, even if found in BF)
    cee_found = False
    for search_path in search_paths:
        cee_index = search_path / "CEE-index.txt"
        if cee_index.exists():
            cee_dict = get_index_rtf_cached(cee_index)
            cee_found = True
            # First try exact match (already normalized)
            if norm_title in cee_dict:
                code_nums = cee_dict[norm_title]  # This is now a list of codes like ["5.I"]
                for code_num in code_nums:
                    code = f"CEE {code_num}"
                    all_matches.append((code, 100.0, norm_title))
            # Always also try fuzzy matching to catch all possible matches
            # Try matching on all entries with dynamic character difference tolerance (ignoring spaces/non-letters)
            for dict_title, code_nums in cee_dict.items():
                # Skip if we already found exact match for this title
                if norm_title == dict_title:
                    continue
                # Both are already normalized, so we can compare directly
                matches, distance, letters_1, letters_2 = matches_with_tolerance(norm_title, dict_title, max_diff)
                
                if matches:
                    # Calculate similarity percentage
                    max_len = max(len(letters_1), len(letters_2))
                    similarity = (1.0 - distance / max_len) * 100.0 if max_len > 0 else 100.0
                    
                    # Add all codes for this matching title
                    for code_num in code_nums:
                        code = f"CEE {code_num}"
                        all_matches.append((code, similarity, dict_title))
            break  # Only check first found file
    
    search_end = time.perf_counter()
    search_time = search_end - search_start
    
    if not all_matches:
        return ""
    
    # Remove duplicates (same code)
    unique_matches = {}
    for code, similarity, matched_title in all_matches:
        if code not in unique_matches or similarity > unique_matches[code][1]:
            unique_matches[code] = (code, similarity, matched_title)
    
    codes = list(unique_matches.keys())
    
    # Load code-to-name mapping once (works for single and multiple matches)
    coduri_dict = get_coduri_cached()
    
    def format_code_with_name(code: str) -> str:
        """
        Return the full line from Coduri-BF-CEE if available.
        Falls back to progressively simpler code forms before returning the raw code.
        """
        # Exact match first (e.g., "BF 5.I")
        if code in coduri_dict:
            return coduri_dict[code]
        
        # Remove trailing "subtip" notation, if present (e.g., "130.I subtip" -> "130.I")
        code_no_subtip = re.sub(r'\s+subtip.*$', '', code, flags=re.IGNORECASE)
        if code_no_subtip in coduri_dict:
            return coduri_dict[code_no_subtip]
        
        # Remove the last dot segment (e.g., "131.II" -> "131")
        base_code = re.sub(r'\.[A-Za-z0-9]+$', '', code_no_subtip)
        if base_code in coduri_dict:
            return coduri_dict[base_code]
        
        # Fallback: return the raw code if no mapping found
        return code
    
    # If multiple matches, include names for each code and keep the " - qqq" suffix
    if len(codes) > 1:
        formatted_codes = [format_code_with_name(code) for code in codes]
        codes_str = ", ".join(formatted_codes)
        return f"{codes_str} - qqq"
    
    # If single match, use the formatted name if available; otherwise keep existing fallback
    single_code = codes[0]  # This is now "BF 5.I" or "CEE 5.I" format
    formatted_single = format_code_with_name(single_code)
    if formatted_single != single_code:
        return formatted_single
    
    # If not found, return the code with "- qqq" suffix
    return f"{single_code} - qqq"


# Cache for index RTF files
_index_rtf_cache = {}


def get_index_rtf_cached(index_file: Path) -> Dict[str, List[str]]:
    """Get index RTF with caching."""
    index_file_str = str(index_file)
    if index_file_str not in _index_rtf_cache:
        _index_rtf_cache[index_file_str] = load_index_rtf(index_file)
    return _index_rtf_cache[index_file_str]


# ---------- OCR Loading Functions ----------

def load_ocr_pages(ocr_root: Path, selected_volumes: list = None):
    """
    Din toate *.txt din folderele selectate extrage paginile OCR.
    Returnează listă de dict cu informații despre pagini.
    selected_volumes: listă de căi către folderele finale (volumes) care conțin ocr.txt
    """
    pages = []

    # Dacă nu sunt volume selectate, nu căuta nimic
    if not selected_volumes:
        return pages

    # Use thread-safe list for parallel loading
    pages_lock = threading.Lock()
    volume_times = {}
    
    def load_volume(volume_path_str: str):
        """Load pages from a single volume (thread-safe)."""
        vol_start = time.perf_counter()
        volume_pages = []
        volume_path = Path(volume_path_str)
        if volume_path.exists() and folder_has_ocr(volume_path):
            _load_pages_from_folder(volume_path, volume_pages)
        vol_end = time.perf_counter()
        vol_time = vol_end - vol_start
        volume_times[volume_path_str] = (len(volume_pages), vol_time)
        return volume_pages

    # Parallelize loading volumes using ThreadPoolExecutor
    # Use max_workers based on number of volumes, but cap at reasonable limit
    max_workers = min(len(selected_volumes), 16)  # Cap at 16 threads for I/O
    
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        # Submit all volume loading tasks
        future_to_volume = {
            executor.submit(load_volume, volume_path_str): volume_path_str
            for volume_path_str in selected_volumes
        }
        
        # Collect results as they complete
        for future in as_completed(future_to_volume):
            try:
                volume_pages = future.result()
                with pages_lock:
                    pages.extend(volume_pages)
            except Exception as e:
                # Log error but continue with other volumes
                print(f"[{time.strftime('%H:%M:%S')}] Error loading volume {future_to_volume[future]}: {e}")

    return pages


def _search_recursive(folder: Path, pages: list):
    """Recursively search for OCR files in folder and subfolders.
    If a folder has OCR, it's an end folder - load from it and don't search subfolders."""
    # Check if this folder has OCR files
    if folder_has_ocr(folder):
        _load_pages_from_folder(folder, pages)
        # Don't search in subfolders if this folder has OCR (it's an end folder)
        return
    
    # No OCR in this folder, search in subfolders
    for subfolder in sorted(p for p in folder.iterdir() if p.is_dir()):
        _search_recursive(subfolder, pages)


def folder_has_ocr(folder: Path) -> bool:
    """Check if a folder contains ocr.txt files (marks it as end folder)."""
    # Check for ocr.txt specifically, or any .txt files
    return (folder / "ocr.txt").exists() or any(folder.glob("*.txt"))


def _load_pages_from_folder(folder: Path, pages: list):
        folder_start = time.perf_counter()
        txt_files = sorted(folder.glob("*.txt"))
        pages_before = len(pages)
        for txt in txt_files:
            file_start = time.perf_counter()
            with txt.open("r", encoding="utf-8") as f:
                current_page_num = None
                current_page_img = None
                buffer = []
                lines_buffer = []
                pages_in_file = 0

                def flush_page():
                    nonlocal buffer, current_page_num, current_page_img, lines_buffer, pages_in_file
                    if current_page_num is None or current_page_img is None:
                        return
                    content = "".join(buffer).strip()
                    norm = normalize_text(content)
                    word_count = len(norm.split())
                    if word_count < MIN_PAGE_WORDS:
                        return
                    clean_lines = [line.strip() for line in lines_buffer if line.strip() and not PAGE_HEADER_RE.match(line.strip())]
                    page_num_from_img = extract_page_number_from_image(current_page_img)
                    
                    # Check if image file exists
                    image_path = folder / current_page_img
                    image_exists = image_path.exists()
                    if not image_exists:
                        # Try to find it in subdirectories
                        for img_file in folder.rglob(current_page_img):
                            if img_file.exists():
                                image_exists = True
                                break
                    
                    pages.append({
                        "folder": folder.name,
                        "volume_path": str(folder),  # Full path to volume folder
                        "txt_file": txt.name,
                        "page_num": page_num_from_img if page_num_from_img else current_page_num,
                        "page_img": current_page_img,
                        "image_exists": image_exists,  # Store whether image exists
                        "text": content,
                        "lines": clean_lines,
                    })
                    pages_in_file += 1

                for line in f:
                    m = PAGE_HEADER_RE.match(line)
                    if m:
                        flush_page()
                        current_page_num = m.group(1)
                        current_page_img = m.group(2)
                        buffer = []
                        lines_buffer = []
                    else:
                        buffer.append(line)
                        lines_buffer.append(line)
                flush_page()
            file_end = time.perf_counter()
            file_time = file_end - file_start
        folder_end = time.perf_counter()
        folder_time = folder_end - folder_start
        pages_added = len(pages) - pages_before


# ---------- Search Functions ----------

def find_match_context(query_text: str, page_lines: list, before: int = 2, after: int = 2):
    """
    Găsește poziția match-ului în linii și returnează contextul.
    """
    norm_query = normalize_text(query_text)
    query_words = norm_query.split()
    
    best_line_idx = None
    best_match_score = 0
    
    for i, line in enumerate(page_lines):
        norm_line = normalize_text(line)
        if not norm_line:
            continue
        norm_line_words = norm_line.split()
        # Ensure each query word matches a different text word
        matched_word_indices = set()
        words_found = 0
        for query_word in query_words:
            for idx, text_word in enumerate(norm_line_words):
                if idx not in matched_word_indices and words_match(query_word, text_word):
                    matched_word_indices.add(idx)
                    words_found += 1
                    break
        if words_found > 0:
            score = fuzz.partial_ratio(norm_query, norm_line)
            if score > best_match_score:
                best_match_score = score
                best_line_idx = i
    
    if best_line_idx is None:
        return page_lines[:before + after + 1] if page_lines else []
    
    start_idx = max(0, best_line_idx - before)
    end_idx = min(len(page_lines), best_line_idx + after + 1)
    
    return page_lines[start_idx:end_idx]


def _search_single_page_wrapper(args):
    """Unpack arguments for process pool worker."""
    return _search_single_page(*args)


def check_words_in_word_span(query_words: list, page_lines: list, word_span: int, exact_order: bool = True, threshold: int = DEFAULT_THRESHOLD) -> tuple:
    """
    Verifică dacă toate cuvintele din query apar într-un span de cuvinte.
    exact_order: True = cuvintele trebuie să apară în ordinea exactă, False = orice ordine
    Returnează (found, score, start_line_idx, end_line_idx, matched_words).
    matched_words: lista cu cuvintele exacte din fragmentul care s-a potrivit (din liniile originale, nu normalizate)
    """
    if not page_lines or not query_words:
        return (False, 0, 0, 0, [])
    
    # Keep original lines for the matched fragment
    original_lines = [line.strip() for line in page_lines if line.strip()]
    norm_lines = [normalize_text(line) for line in original_lines]
    if not norm_lines:
        return (False, 0, 0, 0, [])
    
    # Convert lines to a list of words with their line indices and original word text
    all_words = []  # (normalized_word, line_idx, original_word, original_line)
    for line_idx, (norm_line, orig_line) in enumerate(zip(norm_lines, original_lines)):
        norm_words = norm_line.split()
        # For original words, we need to handle cases where normalization splits words
        # (e.g., "Corbea/ncălica" becomes "corbea ncalica")
        # Split original line and try to map normalized words back
        orig_words_raw = orig_line.split()
        orig_line_norm = normalize_text(orig_line)
        orig_words_norm = orig_line_norm.split()
        
        # Map each normalized word to an original word
        for i, norm_word in enumerate(norm_words):
            orig_word = norm_word  # Default fallback
            # Try to find matching original word
            if i < len(orig_words_norm):
                # Find which original word(s) this normalized word came from
                for orig_w in orig_words_raw:
                    orig_w_norm = normalize_text(orig_w)
                    # Check if this normalized word is in the normalized original word
                    if norm_word == orig_w_norm or (norm_word in orig_w_norm and len(norm_word) >= 4):
                        orig_word = orig_w
                        break
                    # Also check if original word contains this normalized word (for split cases)
                    orig_w_parts = orig_w_norm.split()
                    if norm_word in orig_w_parts:
                        orig_word = orig_w
                        break
            all_words.append((norm_word, line_idx, orig_word, orig_line))
    
    if not all_words:
        return (False, 0, 0, 0, [])
    
    best_score = 0
    best_start_line = 0
    best_end_line = 0
    best_matched_words = []
    found_all = False
    
    # Slide a window of word_span words across the text
    for start_word_idx in range(len(all_words)):
        end_word_idx = min(start_word_idx + word_span, len(all_words))
        span_words_data = all_words[start_word_idx:end_word_idx]
        span_words = [w[0] for w in span_words_data]  # normalized words
        span_text = " ".join(span_words)
        
        # Get line indices for the span
        if span_words_data:
            start_line_idx = span_words_data[0][1]
            end_line_idx = span_words_data[-1][1] + 1
        else:
            start_line_idx = 0
            end_line_idx = 0
        
        # Check word order if exact_order is True
        if exact_order:
            # Check if words appear in exact order
            # Track which span word indices have been matched to ensure each query word matches a different text word
            matched_indices = set()
            word_indices = []
            for query_word in query_words:
                # Find matching word in span_words using fuzzy matching
                found_idx = None
                for idx, span_word in enumerate(span_words):
                    # Only match if this word hasn't been matched yet and it matches the query word
                    if idx not in matched_indices and words_match(query_word, span_word):
                        found_idx = idx
                        matched_indices.add(idx)  # Mark this word as used
                        break
                
                if found_idx is not None:
                    word_indices.append(found_idx)
                else:
                    # Word not found
                    word_indices = None
                    break
            
            if word_indices and len(word_indices) == len(query_words):
                # Check if indices are in ascending order (exact order)
                if all(word_indices[i] <= word_indices[i+1] for i in range(len(word_indices)-1)):
                    words_found = len(query_words)
                else:
                    words_found = 0
            else:
                words_found = 0
        else:
            # Random order: check if all words are present, but each query word must match a different text word
            matched_indices = set()
            words_found = 0
            for query_word in query_words:
                found = False
                for idx, span_word in enumerate(span_words):
                    # Only match if this word hasn't been matched yet and it matches the query word
                    if idx not in matched_indices and words_match(query_word, span_word):
                        matched_indices.add(idx)  # Mark this word as used
                        words_found += 1
                        found = True
                        break
                if not found:
                    # This query word couldn't be matched to a unique text word
                    break
        
        if words_found == 0:
            continue
        
        word_coverage = (words_found / len(query_words)) * 100 if query_words else 0
        query_text = " ".join(query_words)
        fuzz_score = fuzz.partial_ratio(query_text, span_text, score_cutoff=threshold)
        
        if words_found == len(query_words):
            score = max(fuzz_score, word_coverage * 0.9)
            if score > best_score:
                best_score = score
                best_start_line = start_line_idx
                best_end_line = end_line_idx
                # Store the exact matched words (original, not normalized)
                best_matched_words = [w[2] for w in span_words_data]
                found_all = True
        elif words_found >= len(query_words) * 0.8:
            score = (fuzz_score + word_coverage) / 2
            if score > best_score:
                best_score = score
                best_start_line = start_line_idx
                best_end_line = end_line_idx
                best_matched_words = [w[2] for w in span_words_data]
        else:
            score = fuzz_score * (words_found / len(query_words))
            if score > best_score and not found_all:
                best_score = score
                best_start_line = start_line_idx
                best_end_line = end_line_idx
                best_matched_words = [w[2] for w in span_words_data]
    
    return (found_all, best_score, best_start_line, best_end_line, best_matched_words)


def _search_single_page(page, norm_query, query_words, query_text, threshold, word_span, exact_order):
    """
    Search a single page for matches.
    """
    page_matches = []
    page_lines = page.get("lines", [])
    
    if word_span is not None and word_span > 0 and page_lines:
        found, score, start_line_idx, end_line_idx, matched_words = check_words_in_word_span(
            query_words, page_lines, word_span, exact_order, threshold
        )
        
        if found and score >= threshold:
            # Use the exact matched fragment
            matched_fragment = " ".join(matched_words) if matched_words else ""
            
            # Find title for this page
            volume_path = Path(page.get("volume_path", ""))
            title = find_title_for_page(
                volume_path,
                page["page_num"],
                page.get("text", ""),
                start_line_idx,
                page_lines
            )
            
            # Find code for the title
            code = find_code_for_title(volume_path, title)
            
            page_matches.append({
                "folder": page["folder"],
                "volume_path": page.get("volume_path", ""),
                "image": page["page_img"],
                "score": round(score, 1),
                "page_num": page["page_num"],
                "title": title,  # Add title
                "code": code,  # Add code
                "snippet": [matched_fragment],  # Store as list for consistency
                "matched_fragment": matched_fragment,  # Exact matched fragment
                "page_lines": page_lines,  # Store full page lines for context
                "match_start_line_idx": start_line_idx,  # Line index where match starts
                "match_end_line_idx": end_line_idx  # Line index where match ends
            })
        return page_matches
    
    norm_page = normalize_text(page["text"])
    if not norm_page:
        return page_matches
    
    score = fuzz.partial_ratio(norm_query, norm_page, score_cutoff=threshold)
    norm_page_words = norm_page.split()
    # Ensure each query word matches a different text word
    matched_word_indices = set()
    words_found = 0
    for query_word in query_words:
        for idx, text_word in enumerate(norm_page_words):
            if idx not in matched_word_indices and words_match(query_word, text_word):
                matched_word_indices.add(idx)
                words_found += 1
                break
    word_coverage = (words_found / len(query_words)) * 100 if query_words else 0
    
    if words_found == len(query_words):
        final_score = max(score, word_coverage * 0.9)
    elif words_found >= len(query_words) * 0.8:
        final_score = (score + word_coverage) / 2
    else:
        final_score = score * (words_found / len(query_words))
    
    if final_score >= threshold:
        snippet = find_match_context(query_text, page_lines)
        # For full page search, extract the best matching fragment
        # Find the line with the best match and extract words around it
        snippet_text = " ".join(snippet)
        snippet_words = snippet_text.split()
        
        # Find where search words appear in snippet
        norm_snippet_words = [normalize_text(w) for w in snippet_words]
        norm_query_words = [normalize_text(w) for w in query_words]
        
        match_start_idx = None
        for i, norm_word in enumerate(norm_snippet_words):
            for norm_query_word in norm_query_words:
                if words_match(norm_query_word, norm_word):
                    match_start_idx = i
                    break
            if match_start_idx is not None:
                break
        
        # Find the line index in page_lines where the match occurs
        best_line_idx = None
        best_match_score = 0
        for i, line in enumerate(page_lines):
            norm_line = normalize_text(line)
            if not norm_line:
                continue
            norm_line_words = norm_line.split()
            # Ensure each query word matches a different text word
            matched_word_indices = set()
            words_found = 0
            for query_word in query_words:
                for idx, text_word in enumerate(norm_line_words):
                    if idx not in matched_word_indices and words_match(query_word, text_word):
                        matched_word_indices.add(idx)
                        words_found += 1
                        break
            if words_found > 0:
                score = fuzz.partial_ratio(norm_query, norm_line, score_cutoff=threshold)
                if score > best_match_score:
                    best_match_score = score
                    best_line_idx = i
        
        if match_start_idx is not None:
            # Extract a window around the match (use word_span if available, otherwise 20 words)
            fragment_size = word_span if word_span else 20
            half_size = fragment_size // 2
            start_idx = max(0, match_start_idx - half_size)
            end_idx = min(len(snippet_words), start_idx + fragment_size)
            if end_idx - start_idx < fragment_size and start_idx > 0:
                start_idx = max(0, end_idx - fragment_size)
            matched_fragment = " ".join(snippet_words[start_idx:end_idx])
        else:
            # Fallback: use first fragment_size words
            fragment_size = word_span if word_span else 20
            matched_fragment = " ".join(snippet_words[:fragment_size])
        
            # Find title for this page
            volume_path = Path(page.get("volume_path", ""))
            title = find_title_for_page(
                volume_path,
                page["page_num"],
                page.get("text", ""),
                best_line_idx if best_line_idx is not None else 0,
                page_lines
            )
            
            # Find code for the title
            code = find_code_for_title(volume_path, title)
            
            page_matches.append({
                "folder": page["folder"],
                "volume_path": page.get("volume_path", ""),
                "image": page["page_img"],
                "score": round(final_score, 1),
                "page_num": page["page_num"],
                "title": title,  # Add title
                "code": code,  # Add code
                "snippet": snippet,
                "matched_fragment": matched_fragment,  # Exact matched fragment
                "page_lines": page_lines,  # Store full page lines for context
                "match_start_line_idx": best_line_idx if best_line_idx is not None else 0,  # Line index where match occurs
                "match_end_line_idx": best_line_idx + 1 if best_line_idx is not None else 1  # Line index where match ends
            })
    
    return page_matches


def search_text_in_pages(query_text: str, pages, threshold: int = DEFAULT_THRESHOLD, word_span: int = None, exact_order: bool = True, use_inflections: bool = True):
    """
    Caută query_text în toate paginile și returnează toate match-urile cu score >= threshold.
    word_span: numărul de cuvinte în care să caute (None = întreaga pagină)
    exact_order: True = cuvintele trebuie să apară în ordinea exactă, False = orice ordine
    use_inflections: Dacă True, extinde căutarea cu forme flexionare din DEX online
    """
    search_start = time.perf_counter()
    print(f"[{time.strftime('%H:%M:%S')}] Starting search in {len(pages)} pages...")
    
    # Cache pentru forme flexionare (partajat între apeluri)
    if not hasattr(search_text_in_pages, 'inflection_cache'):
        search_text_in_pages.inflection_cache = {}
    
    # Extinde cuvintele cu forme flexionare dacă este activat
    inflection_start = time.perf_counter()
    if use_inflections:
        expanded_words = expand_search_terms_with_inflections(
            query_text, 
            search_text_in_pages.inflection_cache
        )
        # Creează un query extins pentru normalizare
        expanded_query = " ".join(expanded_words)
        norm_query = normalize_text(expanded_query)
        query_words = norm_query.split()
    else:
        norm_query = normalize_text(query_text)
        query_words = norm_query.split()
    inflection_end = time.perf_counter()
    inflection_time = inflection_end - inflection_start
    print(f"[{time.strftime('%H:%M:%S')}] Query processing (inflections): {inflection_time:.2f}s, {len(query_words)} words")
    
    if len(query_words) < 1:
        raise ValueError("Query must have at least 1 word")
    if len(query_words) > 10:
        raise ValueError("Query must have at most 10 words")
    
    matches = []
    
    search_exec_start = time.perf_counter()
    max_workers = min(os.cpu_count() or 4, 8)
    print(f"[{time.strftime('%H:%M:%S')}] Using process pool search with {max_workers} workers")
    
    with ProcessPoolExecutor(max_workers=max_workers) as executor:
        future_to_idx = {
            executor.submit(
                _search_single_page_wrapper,
                (page, norm_query, query_words, query_text, threshold, word_span, exact_order)
            ): idx
            for idx, page in enumerate(pages)
        }
        
        for completed_idx, future in enumerate(as_completed(future_to_idx), 1):
            try:
                page_matches = future.result()
                if page_matches:
                    matches.extend(page_matches)
            except Exception as e:
                print(f"[{time.strftime('%H:%M:%S')}] Error searching page: {e}")
            
            if completed_idx % 100 == 0:
                print(f"[{time.strftime('%H:%M:%S')}]   Searched {completed_idx}/{len(pages)} pages...")
    
    search_exec_end = time.perf_counter()
    search_exec_time = search_exec_end - search_exec_start
    print(f"[{time.strftime('%H:%M:%S')}] Search execution completed: {len(matches)} matches found in {search_exec_time:.2f}s")
    
    sort_start = time.perf_counter()
    matches.sort(key=lambda x: x["score"], reverse=True)
    sort_end = time.perf_counter()
    sort_time = sort_end - sort_start
    print(f"[{time.strftime('%H:%M:%S')}] Sorting matches: {sort_time:.2f}s")
    
    search_end = time.perf_counter()
    total_search_time = search_end - search_start
    print(f"[{time.strftime('%H:%M:%S')}] Total search time: {total_search_time:.2f}s")
    
    return matches


def run_search_only(query_text: str, ocr_root: Path, word_span: int = None, threshold: int = DEFAULT_THRESHOLD, selected_volumes: list = None, exact_order: bool = True):
    """
    Run the search and return matches without generating documents.
    selected_volumes: list of paths to volumes (folders with ocr.txt) to search
    word_span: number of words to search within (None = entire page)
    exact_order: True = words must appear in exact order, False = any order
    Returns: list of match dictionaries
    """
    overall_start = time.perf_counter()
    print(f"\n{'='*60}")
    print(f"[{time.strftime('%H:%M:%S')}] ===== STARTING SEARCH =====")
    print(f"[{time.strftime('%H:%M:%S')}] Query: '{query_text}'")
    print(f"[{time.strftime('%H:%M:%S')}] Volumes: {len(selected_volumes) if selected_volumes else 0}")
    print(f"[{time.strftime('%H:%M:%S')}] Word span: {word_span}, Exact order: {exact_order}, Threshold: {threshold}")
    print(f"{'='*60}\n")
    
    if not ocr_root.is_dir():
        raise Exception(f"OCR root folder not found: {ocr_root}")
    
    if threshold < 0 or threshold > 100:
        raise Exception(f"Threshold must be between 0 and 100, got: {threshold}")
    
    # Load OCR pages (only from selected volumes)
    load_start = time.perf_counter()
    pages = load_ocr_pages(ocr_root, selected_volumes)
    load_end = time.perf_counter()
    load_time = load_end - load_start
    print(f"[{time.strftime('%H:%M:%S')}] === Loading phase: {load_time:.2f}s ===\n")
    
    if not pages:
        raise Exception("No pages found in selected folders. Please check your selection.")
    
    # Search
    search_start = time.perf_counter()
    matches = search_text_in_pages(query_text, pages, threshold, word_span=word_span, exact_order=exact_order)
    search_end = time.perf_counter()
    search_time = search_end - search_start
    print(f"[{time.strftime('%H:%M:%S')}] === Search phase: {search_time:.2f}s ===\n")
    
    overall_end = time.perf_counter()
    overall_time = overall_end - overall_start
    print(f"{'='*60}")
    print(f"[{time.strftime('%H:%M:%S')}] ===== SEARCH COMPLETE =====")
    print(f"[{time.strftime('%H:%M:%S')}] Total time: {overall_time:.2f}s")
    print(f"[{time.strftime('%H:%M:%S')}] Breakdown:")
    print(f"[{time.strftime('%H:%M:%S')}]   - Loading: {load_time:.2f}s ({load_time/overall_time*100:.1f}%)")
    print(f"[{time.strftime('%H:%M:%S')}]   - Searching: {search_time:.2f}s ({search_time/overall_time*100:.1f}%)")
    print(f"[{time.strftime('%H:%M:%S')}] Results: {len(matches)} matches found")
    print(f"{'='*60}\n")
    
    return matches


def run_search_and_generate_report(query_text: str, ocr_root: Path, output_dir: Path, word_span: int = None, threshold: int = DEFAULT_THRESHOLD, selected_volumes: list = None, exact_order: bool = True):
    """
    Run the search and generate the report document.
    This replaces the subprocess call to search_text.py
    NOTE: This function is kept for future use but is not currently called.
    selected_volumes: list of paths to volumes (folders with ocr.txt) to search
    word_span: number of words to search within (None = entire page)
    exact_order: True = words must appear in exact order, False = any order
    """
    if not ocr_root.is_dir():
        raise Exception(f"OCR root folder not found: {ocr_root}")
    
    if threshold < 0 or threshold > 100:
        raise Exception(f"Threshold must be between 0 and 100, got: {threshold}")
    
    # Load OCR pages (only from selected volumes)
    pages = load_ocr_pages(ocr_root, selected_volumes)
    
    if not pages:
        raise Exception("No pages found in selected folders. Please check your selection.")
    
    # Search
    matches = search_text_in_pages(query_text, pages, threshold, word_span=word_span, exact_order=exact_order)
    
    # Write results to text file
    output_file = output_dir / "search-result.txt"
    with output_file.open("w", encoding="utf-8") as f:
        if not matches:
            f.write(f"No matches found above threshold {threshold}\n")
        else:
            for match in matches:
                for snippet_line in match['snippet']:
                    if snippet_line.strip():
                        f.write(snippet_line + "\n")
                f.write(f"({match['folder']}, p. {match['page_num']})\n")
                f.write("\n")
    
    # Write results to Word document
    docx_file = None
    
    # Check if running as PyInstaller executable and look for bundled default.docx
    if getattr(sys, 'frozen', False) and hasattr(sys, '_MEIPASS'):
        bundled_docx = Path(sys._MEIPASS) / "default.docx"
        if bundled_docx.exists():
            docx_file = bundled_docx
    
    # If not found in bundle, check output_dir
    if docx_file is None or not docx_file.exists():
        docx_file = output_dir / "default.docx"
    
    # Load template if it exists, otherwise create new document
    if docx_file.exists():
        doc = Document(str(docx_file))
        # Clear existing content (keep styles)
        for para in list(doc.paragraphs):
            p = para._element
            p.getparent().remove(p)
    else:
        doc = Document()
    
    # Add title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run('Search results for "')
    search_run = title_para.add_run(query_text)
    search_run.italic = True
    title_para.add_run('"')
    doc.add_paragraph()
    
    # Add matches
    if not matches:
        doc.add_paragraph(f"No matches found above threshold {threshold}")
    else:
        for match in matches:
            for snippet_line in match['snippet']:
                if snippet_line.strip():
                    para = doc.add_paragraph(snippet_line.strip())
                    try:
                        para.style = "2-Versuri-centru"
                    except KeyError:
                        pass
            # Build source text with title if available
            title = match.get("title", "")
            if title:
                source_text = f"({match['folder']}, {title}, p. {match['page_num']})"
            else:
                source_text = f"({match['folder']}, p. {match['page_num']})"
            para = doc.add_paragraph(source_text)
            try:
                para.style = "4-Sursa text"
            except KeyError:
                pass
            doc.add_paragraph()
    
    # Save to output_dir
    doc.save(str(output_dir / "default.docx"))


def open_image(image_path: Path):
    """Open an image file using the system's default image viewer."""
    if not image_path.exists():
        messagebox.showerror("Error", f"Image not found: {image_path}")
        return
    
    try:
        system = platform.system()
        if system == "Darwin":  # macOS
            subprocess.run(["open", str(image_path)])
        elif system == "Windows":
            import os
            os.startfile(str(image_path))
        else:  # Linux and others
            subprocess.run(["xdg-open", str(image_path)])
    except Exception as e:
        messagebox.showerror("Error", f"Failed to open image: {str(e)}")


def _sort_images_by_page(images: list) -> list:
    """Sort images by page number extracted from filename; fallback to name."""
    def sort_key(p: Path):
        num = extract_page_number_from_image(p.name)
        try:
            num_int = int(num) if num else None
        except ValueError:
            num_int = None
        if num_int is not None:
            return (0, num_int, p.name.lower())
        return (1, p.name.lower())
    return sorted(images, key=sort_key)


def _list_neighbor_images(image_path: Path) -> list:
    """List images in the same folder, sorted by page number if present."""
    folder = image_path.parent
    patterns = ["*.png", "*.jpg", "*.jpeg", "*.webp", "*.tif", "*.tiff", "*.bmp"]
    imgs = []
    for pattern in patterns:
        imgs.extend(folder.glob(pattern))
    # Deduplicate while preserving order before sort
    imgs = list(dict.fromkeys(imgs))
    return _sort_images_by_page(imgs)


def _list_neighbor_pdfs(pdf_path: Path) -> list:
    """List PDFs in the same folder, sorted by page number if present."""
    folder = pdf_path.parent
    pdfs = list(folder.glob("*.pdf"))
    pdfs = list(dict.fromkeys(pdfs))
    return _sort_images_by_page(pdfs)


def _get_viewer_image_max_size() -> tuple[int, int]:
    """
    Compute the max width/height available for rendering inside the viewer window.
    Falls back to screen size when the window is not yet sized.
    """
    state = image_viewer_state
    win = state.get("window")
    if win and win.winfo_exists():
        win.update_idletasks()
        # Leave room for paddings/buttons/title
        max_w = max(400, win.winfo_width() - 40)
        max_h = max(400, win.winfo_height() - 160)
    else:
        if root:
            screen_w = root.winfo_screenwidth()
            screen_h = root.winfo_screenheight()
        else:
            screen_w, screen_h = 1600, 900
        max_w = int(screen_w * 0.9)
        max_h = int(screen_h * 0.9)
    return max_w, max_h


def _refresh_current_view():
    """Re-render current image/PDF page to fit the latest window size."""
    state = image_viewer_state
    if state.get("is_pdf"):
        if state.get("pdf_doc"):
            _show_pdf_page(state.get("page_index", 0))
    else:
        if state.get("images"):
            _show_image_in_viewer(state.get("index", 0))


def _on_viewer_resize(event=None):
    """Debounced resize handler to avoid cutting off top/bottom content."""
    state = image_viewer_state
    win = state.get("window")
    if not win or not win.winfo_exists():
        return
    if state.get("resize_job"):
        try:
            win.after_cancel(state["resize_job"])
        except Exception:
            pass
    state["resize_job"] = win.after(120, _refresh_current_view)


def _show_image_in_viewer(index: int):
    """Render the image at the given index inside the viewer window."""
    state = image_viewer_state
    if not state["images"]:
        return
    index = max(0, min(index, len(state["images"]) - 1))
    state["index"] = index
    state["file_index"] = index
    path = state["images"][index]

    # Fallback if Pillow not available: open with OS viewer
    if Image is None or ImageTk is None:
        open_image(path)
        return

    try:
        img = Image.open(path)
        max_w, max_h = _get_viewer_image_max_size()
        try:
            img.thumbnail((max_w, max_h), Image.Resampling.LANCZOS)  # Pillow >=9.1
        except Exception:
            img.thumbnail((max_w, max_h))
        photo = ImageTk.PhotoImage(img)
    except Exception as e:
        messagebox.showerror("Error", f"Failed to load image: {e}")
        return

    state["photo"] = photo  # keep reference
    state["label"].config(image=photo)
    state["title"].config(text=f"{path.name}  ({index + 1}/{len(state['images'])})")


def _show_pdf_page(index: int):
    """Render the PDF page at the current file/page indices inside the viewer."""
    state = image_viewer_state
    if not state.get("is_pdf") or state.get("pdf_doc") is None:
        return
    total = state.get("pdf_total", 0)
    if total <= 0:
        return
    page_idx = max(0, min(state["page_index"], total - 1))
    state["page_index"] = page_idx

    if fitz is None:
        messagebox.showerror("Error", "PyMuPDF (fitz) not installed; cannot preview PDF.")
        return
    if Image is None or ImageTk is None:
        messagebox.showerror("Error", "Pillow not installed; cannot render PDF.")
        return

    try:
        page = state["pdf_doc"].load_page(page_idx)
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # modest zoom
        mode = "RGBA" if pix.alpha else "RGB"
        img = Image.frombytes(mode, (pix.width, pix.height), pix.samples)
        max_w, max_h = _get_viewer_image_max_size()
        try:
            img.thumbnail((max_w, max_h), Image.Resampling.LANCZOS)
        except Exception:
            img.thumbnail((max_w, max_h))
        photo = ImageTk.PhotoImage(img)
    except Exception as e:
        messagebox.showerror("Error", f"Failed to render PDF page: {e}")
        return

    state["photo"] = photo
    state["label"].config(image=photo)
    file_total = len(state.get("file_paths", [])) or 1
    fname = Path(state["file_paths"][state["file_index"]]).name if state.get("file_paths") else "PDF"
    state["title"].config(
        text=f"{fname}  (file {state['file_index'] + 1}/{file_total}, p. {page_idx + 1}/{total})"
    )


def _load_and_show_pdf(file_idx: int, page_idx: int = 0):
    """Load the PDF at file_idx and render the requested page."""
    state = image_viewer_state
    file_paths = state.get("file_paths", [])
    if not file_paths:
        return
    file_idx = max(0, min(file_idx, len(file_paths) - 1))
    pdf_path = file_paths[file_idx]

    # Close previous doc
    try:
        if state.get("pdf_doc"):
            state["pdf_doc"].close()
    except Exception:
        pass

    try:
        doc = fitz.open(pdf_path)
    except Exception as e:
        messagebox.showerror("Error", f"Failed to open PDF: {e}")
        return

    # Determine total pages
    try:
        page_total = getattr(doc, "page_count", None)
        if page_total is None:
            page_total = len(doc)
    except Exception:
        page_total = len(doc)

    if page_total <= 1 and PyPDF2 is not None:
        try:
            reader = PyPDF2.PdfReader(str(pdf_path))
            page_total = len(reader.pages)
        except Exception:
            pass

    state["file_index"] = file_idx
    state["page_index"] = max(0, min(page_idx, page_total - 1))
    state["pdf_doc"] = doc
    state["pdf_total"] = page_total

    _show_pdf_page(state["page_index"])


def _navigate_pdf(delta_pages: int):
    """Navigate pages/files: left/right arrows jump across files when needed."""
    state = image_viewer_state
    if not state.get("is_pdf"):
        return
    file_paths = state.get("file_paths", [])
    if not file_paths:
        return

    new_page = state.get("page_index", 0) + delta_pages
    file_idx = state.get("file_index", 0)

    if state.get("pdf_total", 0) <= 0:
        return

    # Move across files when page goes out of bounds
    while True:
        if new_page < 0:
            file_idx -= 1
            if file_idx < 0:
                file_idx = 0
                new_page = 0
                break
            # load previous file and jump to its last page
            _load_and_show_pdf(file_idx, 0)
            new_page = state["pdf_total"] - 1
        elif new_page >= state.get("pdf_total", 1):
            file_idx += 1
            if file_idx >= len(file_paths):
                file_idx = len(file_paths) - 1
                new_page = state["pdf_total"] - 1
                break
            # load next file and start at first page
            _load_and_show_pdf(file_idx, 0)
            new_page = 0
        else:
            break

    # Finally render the computed position
    if state.get("file_index") != file_idx or state.get("pdf_doc") is None:
        _load_and_show_pdf(file_idx, new_page)
    else:
        state["page_index"] = new_page
        _show_pdf_page(new_page)


def open_pdf_with_navigation(pdf_path: Path):
    """Open a PDF with page-by-page navigation using PyMuPDF + Pillow."""
    # Fallback if dependencies missing
    if fitz is None or Image is None or ImageTk is None:
        open_document(pdf_path)
        return

    # Collect neighbor PDFs
    pdfs = _list_neighbor_pdfs(pdf_path)
    if pdf_path not in pdfs and pdf_path.exists():
        pdfs.append(pdf_path)
        pdfs = _sort_images_by_page(pdfs)

    state = image_viewer_state

    # Close existing viewer window if open
    if state["window"] and state["window"].winfo_exists():
        try:
            state["window"].destroy()
        except Exception:
            pass

    # Initialize viewer state for PDFs
    state.update({
        "window": None,
        "images": [],
        "file_paths": pdfs,
        "file_index": 0,
        "index": 0,
        "page_index": 0,
        "photo": None,
        "label": None,
        "title": None,
        "pdf_doc": None,
        "pdf_total": 0,
        "is_pdf": True,
        "resize_job": None,
    })

    win = tk.Toplevel(root)
    win.title("Preview PDF")
    win.minsize(500, 500)
    state["window"] = win

    def _on_close():
        try:
            if state["pdf_doc"]:
                state["pdf_doc"].close()
        except Exception:
            pass
        win.destroy()

    win.protocol("WM_DELETE_WINDOW", _on_close)

    state["title"] = tk.Label(win, font=("Arial", 16))
    state["title"].pack(pady=4)

    state["label"] = tk.Label(win, bg="black")
    state["label"].pack(padx=10, pady=10, expand=True)

    btn_frame = tk.Frame(win)
    btn_frame.pack(pady=6)

    prev_btn = tk.Button(btn_frame, text="⬅️", font=("Arial", 20),
                         command=lambda: _navigate_pdf(-1))
    next_btn = tk.Button(btn_frame, text="➡️", font=("Arial", 20),
                         command=lambda: _navigate_pdf(1))
    prev_btn.grid(row=0, column=0, padx=8)
    next_btn.grid(row=0, column=1, padx=8)

    win.bind("<Left>", lambda _: _navigate_pdf(-1))
    win.bind("<Right>", lambda _: _navigate_pdf(1))
    win.bind("<Configure>", _on_viewer_resize)
    win.focus_set()

    # Jump to the requested pdf in the list
    start_idx = pdfs.index(pdf_path) if pdf_path in pdfs else 0
    state["file_index"] = start_idx
    state["page_index"] = 0
    _load_and_show_pdf(state["file_index"], state["page_index"])


def open_image_with_navigation(image_path: Path):
    """
    Open an image with left/right navigation between neighboring pages.
    Supports common raster formats; PDFs handled separately.
    """
    imgs = _list_neighbor_images(image_path)
    if image_path not in imgs and image_path.exists():
        imgs.append(image_path)
        imgs = _sort_images_by_page(imgs)

    state = image_viewer_state
    if state["window"] and state["window"].winfo_exists():
        state["window"].destroy()

    # Reset viewer state for images
    state.update({
        "images": imgs,
        "file_paths": imgs,
        "index": 0,
        "file_index": 0,
        "page_index": 0,
        "photo": None,
        "label": None,
        "title": None,
        "pdf_doc": None,
        "pdf_total": 0,
        "is_pdf": False,
        "resize_job": None,
    })
    win = tk.Toplevel(root)
    win.title("Preview pagină")
    win.minsize(500, 500)
    state["window"] = win

    state["title"] = tk.Label(win, font=("Arial", 16))
    state["title"].pack(pady=4)

    state["label"] = tk.Label(win, bg="black")
    state["label"].pack(padx=10, pady=10, expand=True)

    btn_frame = tk.Frame(win)
    btn_frame.pack(pady=6)

    prev_btn = tk.Button(btn_frame, text="⬅️", font=("Arial", 20),
                         command=lambda: _show_image_in_viewer(state["index"] - 1))
    next_btn = tk.Button(btn_frame, text="➡️", font=("Arial", 20),
                         command=lambda: _show_image_in_viewer(state["index"] + 1))
    prev_btn.grid(row=0, column=0, padx=8)
    next_btn.grid(row=0, column=1, padx=8)

    win.bind("<Left>", lambda _: _show_image_in_viewer(state["index"] - 1))
    win.bind("<Right>", lambda _: _show_image_in_viewer(state["index"] + 1))
    win.bind("<Configure>", _on_viewer_resize)
    win.focus_set()

    start_idx = imgs.index(image_path) if image_path in imgs else 0
    _show_image_in_viewer(start_idx)


def open_document(doc_path: Path):
    """Open a document file using the system's default application."""
    if not doc_path.exists():
        messagebox.showerror("Error", f"Document not found: {doc_path}")
        return
    
    try:
        system = platform.system()
        if system == "Darwin":  # macOS
            subprocess.run(["open", str(doc_path)])
        elif system == "Windows":
            import os
            os.startfile(str(doc_path))
        else:  # Linux and others
            subprocess.run(["xdg-open", str(doc_path)])
    except Exception as e:
        messagebox.showerror("Error", f"Failed to open document: {str(e)}")


def format_quote_with_bold(quote_text: str, search_words: list) -> str:
    """
    Format quote text with search words marked for display.
    Returns the text with search words wrapped in markers for bold formatting.
    For now, we'll use a simple approach and mark words with ** for display.
    """
    if not quote_text or not search_words:
        return quote_text
    
    # Normalize search words
    norm_search_words = [normalize_text(word) for word in search_words]
    
    # Split quote into words and mark matches
    words = quote_text.split()
    result_words = []
    
    for word in words:
        norm_word = normalize_text(word)
        # Check if this word matches any search word
        is_match = any(search_word in norm_word or norm_word in search_word for search_word in norm_search_words)
        
        if is_match:
            # Mark for bold (we'll use a special marker that we'll replace in the UI)
            result_words.append(f"**{word}**")
        else:
            result_words.append(word)
    
    return " ".join(result_words)


def find_image_path(volume_path: Path, image_filename: str) -> Path:
    """Find the full path to an image file given the volume path and image filename."""
    # Image should be in the same folder as the OCR txt file
    image_path = volume_path / image_filename
    if image_path.exists():
        return image_path
    
    # Try to find it in subdirectories
    for img_file in volume_path.rglob(image_filename):
        if img_file.exists():
            return img_file
    
    return image_path  # Return even if not found, let open_image handle the error


def on_select_folder():
    """Handle the Select Biblioteca button click."""
    global selected_root_folder, selected_paths, selected_volumes
    global folder_label, library_tree
    
    # Get last library directory or use home directory
    last_dir = get_last_library_dir()
    initial_dir = last_dir if last_dir else str(Path.home())
    
    folder = filedialog.askdirectory(title="Select Biblioteca Folder", initialdir=initial_dir)
    if folder:
        selected_root_folder = folder
        selected_paths = []  # Reset selections
        selected_volumes = []  # Reset volumes
        folder_label.config(text=f"Biblioteca: {Path(folder).name}")
        populate_library_tree(folder)
        # Remember the directory for next time
        save_last_library_dir(folder)
    else:
        folder_label.config(text="No Biblioteca selected")
        library_tree.delete(*library_tree.get_children())
        selected_volumes = []  # Clear volumes


def generate_filename_from_search(search_term: str) -> str:
    """Generate filename from search term: first 3 words, hyphenated, with '...' if more."""
    words = search_term.split()
    if len(words) <= 3:
        filename = "-".join(words)
    else:
        filename = "-".join(words[:3]) + "..."
    # Remove any invalid filename characters
    filename = "".join(c for c in filename if c.isalnum() or c in "-._")
    return filename + ".docx"


def create_results_window(search_term: str = ""):
    """Create a new results window."""
    global results_window, results_scrollable_frame, results_canvas, export_button_frame, results_title_label
    
    # Always create a new window (don't reuse existing ones)
    # Create new results window
    results_window = tk.Toplevel(root)
    if search_term:
        # Truncate search term if too long for title
        title_search = search_term[:50] + "..." if len(search_term) > 50 else search_term
        results_window.title(f"Rezultate Căutare: {title_search}")
    else:
        results_window.title("Rezultate Căutare")
    results_window.minsize(800, 500)
    results_window.configure(bg=COLOR_BACKGROUND)
    
    # Ensure window appears in front, especially on Windows when main window is fullscreen
    results_window.lift()
    results_window.focus_force()
    # Temporarily set topmost to ensure it appears above fullscreen windows
    results_window.attributes('-topmost', True)
    results_window.update()
    results_window.attributes('-topmost', False)
    
    # Header
    header_frame = tk.Frame(results_window, bg=COLOR_BACKGROUND, pady=10)
    header_frame.pack(fill=tk.X, padx=15)
    
    results_title_label = tk.Text(
        header_frame,
        font=("Arial", 36, "bold"),
        bg=COLOR_BACKGROUND,
        fg=COLOR_TEXT,
        height=1,
        wrap=tk.WORD,
        relief=tk.FLAT,
        padx=0,
        pady=0
    )
    results_title_label.pack(anchor="w")
    results_title_label.insert("1.0", "📋 Rezultate Căutare")
    results_title_label.config(state=tk.DISABLED)
    
    # Results container
    results_container = tk.Frame(results_window, bg=COLOR_BACKGROUND)
    results_container.pack(fill=tk.BOTH, expand=True, padx=15, pady=10)
    
    # Create scrollable frame for results table
    results_canvas = tk.Canvas(results_container, bg=COLOR_BACKGROUND, highlightthickness=0)
    results_scrollbar = ttk.Scrollbar(results_container, orient="vertical", command=results_canvas.yview)
    results_scrollable_frame = tk.Frame(results_canvas, bg=COLOR_BACKGROUND)
    
    results_scrollable_frame.bind(
        "<Configure>",
        lambda e: results_canvas.configure(scrollregion=results_canvas.bbox("all"))
    )
    
    results_canvas.create_window((0, 0), window=results_scrollable_frame, anchor="nw")
    results_canvas.configure(yscrollcommand=results_scrollbar.set)
    
    results_canvas.pack(side="left", fill="both", expand=True)
    results_scrollbar.pack(side="right", fill="y")
    
    # Bind mousewheel to canvas
    def _on_mousewheel(event):
        results_canvas.yview_scroll(int(-1*(event.delta/120)), "units")
    results_canvas.bind_all("<MouseWheel>", _on_mousewheel)
    
    # Update canvas width when scrollable frame changes
    def configure_canvas_width(event):
        canvas_width = event.width
        results_canvas.itemconfig(results_canvas.find_all()[0], width=canvas_width)
    
    results_scrollable_frame.bind("<Configure>", lambda e: results_canvas.configure(scrollregion=results_canvas.bbox("all")))
    results_canvas.bind("<Configure>", configure_canvas_width)
    
    # Export button frame at the bottom
    export_button_frame = tk.Frame(results_window, bg=COLOR_BACKGROUND, pady=10)
    export_button_frame.pack(fill=tk.X, padx=15, pady=(0, 10))


def display_results_in_table(matches: list, search_term: str, word_span: int = 10):
    """Display search results in a table format.
    word_span: number of words to show in the quote fragment."""
    global results_title_label
    
    # Create new results window with search term
    create_results_window(search_term)
    
    # Update title with result count and search terms
    num_results = len(matches)
    search_words = search_term.split()
    
    # Enable text widget for editing
    results_title_label.config(state=tk.NORMAL)
    results_title_label.delete("1.0", tk.END)
    
    # Insert base text
    results_title_label.insert("1.0", "📋 Rezultate Căutare ")
    
    # Insert search terms in bold
    for i, word in enumerate(search_words):
        # Get position before inserting word
        start_pos = results_title_label.index(tk.END)
        # Insert the word
        results_title_label.insert(tk.END, word)
        # Get position after inserting word
        end_pos = results_title_label.index(tk.END)
        # Apply bold formatting to the word
        results_title_label.tag_add("bold", start_pos, end_pos)
        
        # Add comma and space if not last word
        if i < len(search_words) - 1:
            results_title_label.insert(tk.END, ", ")
    
    # Insert count
    results_title_label.insert(tk.END, f" ({num_results} citate)")
    
    # Configure bold tag
    results_title_label.tag_config("bold", font=("Arial", 36, "bold"))
    
    # Make read-only
    results_title_label.config(state=tk.DISABLED)
    
    # Clear existing results
    for widget in results_scrollable_frame.winfo_children():
        widget.destroy()
    
    if not matches:
        no_results_label = tk.Label(
            results_scrollable_frame,
            text="Nu s-au găsit rezultate.",
            font=("Arial", 20),
            bg=COLOR_BACKGROUND,
            fg=COLOR_TEXT
        )
        no_results_label.pack(pady=10)
        return
    
    # Clear match checkboxes list
    global match_checkboxes
    match_checkboxes = []
    
    # Get search words for highlighting
    search_words = search_term.split()
    
    # Create rows for each match
    for match in matches:
        row_frame = tk.Frame(results_scrollable_frame, bg=COLOR_BACKGROUND, relief=tk.RIDGE, bd=1)
        row_frame.pack(fill=tk.X, pady=2, padx=5)
        
        # Checkbox (checked by default) - custom larger checkbox
        checkbox_var = tk.BooleanVar(value=True)
        
        checkbox_btn = tk.Button(
            row_frame,
            text="☑",
            font=("Arial", 32),
            bg=COLOR_BACKGROUND,
            fg=COLOR_TEXT,
            relief=tk.FLAT,
            bd=0,
            padx=5,
            pady=5,
            cursor="hand2"
        )
        
        def update_checkbox_display(var=None, index=None, mode=None, btn=checkbox_btn, var_ref=checkbox_var):
            if var_ref.get():
                btn.config(text="☑", font=("Arial", 32))
            else:
                btn.config(text="☐", font=("Arial", 32))
        
        def toggle_checkbox(var_ref=checkbox_var):
            var_ref.set(not var_ref.get())
        
        checkbox_btn.config(command=toggle_checkbox)
        checkbox_btn.pack(side=tk.LEFT, padx=5, pady=5)
        match_checkboxes.append((match, checkbox_var))
        
        # Update display when variable changes
        checkbox_var.trace_add("write", update_checkbox_display)
        
        # Create a frame for volume and title (stacked vertically)
        volume_title_frame = tk.Frame(row_frame, bg=COLOR_BACKGROUND)
        volume_title_frame.pack(side=tk.LEFT, padx=5, pady=5)
        
        # Volume name
        volume_label = tk.Label(
            volume_title_frame,
            text=match["folder"],
            font=("Arial", 20),
            bg=COLOR_BACKGROUND,
            fg=COLOR_TEXT,
            anchor="w",
            wraplength=200
        )
        volume_label.pack(anchor="w")
        
        # Title (if available) - displayed under volume name
        title = match.get("title", "")
        if title:
            title_label = tk.Label(
                volume_title_frame,
                text=title,
                font=("Arial", 18, "italic"),
                bg=COLOR_BACKGROUND,
                fg=COLOR_ACCENT,
                anchor="w",
                wraplength=300
            )
            title_label.pack(anchor="w", pady=(2, 0))
            
            # Code (if available) - displayed under title
            code = match.get("code", "")
            if code:
                code_label = tk.Label(
                    volume_title_frame,
                    text=code,
                    font=("Arial", 16),
                    bg=COLOR_BACKGROUND,
                    fg=COLOR_TEXT,
                    anchor="w",
                    wraplength=300
                )
                code_label.pack(anchor="w", pady=(2, 0))
        
        # Check if image exists (fallback: don't show page or button if image doesn't exist)
        image_exists = match.get("image_exists", True)  # Default to True for backward compatibility
        
        if image_exists:
            # Open image button (moved between Volume and Page)
            def make_open_button(match_data):
                btn = tk.Button(
                    row_frame,
                    text="📷",
                    font=("Arial", 18),
                    bg=COLOR_BOOK,
                    fg=COLOR_TEXT,
                    relief=tk.RAISED,
                    padx=8,
                    pady=5,
                    cursor="hand2",
                    command=lambda: open_match_image(match_data)
                )
                return btn
            
            open_btn = make_open_button(match)
            open_btn.pack(side=tk.LEFT, padx=5, pady=5)
            
            # Page number
            page_label = tk.Label(
                row_frame,
                text=f"p. {match['page_num']}",
                font=("Arial", 20),
                bg=COLOR_BACKGROUND,
                fg=COLOR_TEXT
            )
            page_label.pack(side=tk.LEFT, padx=5, pady=5)
        else:
            # Image doesn't exist - skip page number and button
            pass
        
        # Quote with bolded search words - use the exact matched fragment if available
        if "matched_fragment" in match and match["matched_fragment"]:
            # Use the exact matched fragment that was found during search
            matched_fragment = match["matched_fragment"]
            
            # Add 4 words before and 4 words after the matched fragment
            if "page_lines" in match and match["page_lines"]:
                # Get all words from page lines around the match area
                page_lines = match["page_lines"]
                match_start_line = match.get("match_start_line_idx", 0)
                match_end_line = match.get("match_end_line_idx", match_start_line + 1)
                
                # Get context lines (a few lines before and after the match)
                context_start = max(0, match_start_line - 2)
                context_end = min(len(page_lines), match_end_line + 2)
                context_lines = page_lines[context_start:context_end]
                
                # Convert context lines to words
                context_text = " ".join(context_lines)
                all_words = context_text.split()
                matched_words = matched_fragment.split()
                
                # Find where the matched fragment appears in all_words (using normalized comparison)
                matched_start_idx = None
                norm_matched_words = [normalize_text(w) for w in matched_words]
                for i in range(len(all_words) - len(matched_words) + 1):
                    # Compare normalized versions
                    norm_window = [normalize_text(w) for w in all_words[i:i+len(matched_words)]]
                    if norm_window == norm_matched_words:
                        matched_start_idx = i
                        break
                
                if matched_start_idx is not None:
                    # Extract 4 words before and 4 words after
                    start_idx = max(0, matched_start_idx - 4)
                    end_idx = min(len(all_words), matched_start_idx + len(matched_words) + 4)
                    quote_text = " ".join(all_words[start_idx:end_idx])
                else:
                    # Fallback: use matched fragment if we can't find it in context
                    quote_text = matched_fragment
            else:
                # Fallback: use matched fragment if no page_lines available
                quote_text = matched_fragment
        else:
            # Fallback: extract from snippet (for backward compatibility)
            full_quote_text = " ".join(match["snippet"])
            quote_words = full_quote_text.split()
            
            # Find where the search words appear in the quote
            norm_quote_words = [normalize_text(w) for w in quote_words]
            norm_search_words = [normalize_text(w) for w in search_words]
            
            # Find the first occurrence of any search word
            match_start_idx = None
            for i, norm_word in enumerate(norm_quote_words):
                for norm_search_word in norm_search_words:
                    if norm_search_word in norm_word or norm_word in norm_search_word:
                        match_start_idx = i
                        break
                if match_start_idx is not None:
                    break
            
            # If no match found, use the beginning
            if match_start_idx is None:
                match_start_idx = 0
            
            # Calculate window around the match
            # Try to center the match in the window
            half_span = word_span // 2
            start_idx = max(0, match_start_idx - half_span)
            end_idx = min(len(quote_words), start_idx + word_span)
            
            # Adjust start if we're near the end
            if end_idx - start_idx < word_span and start_idx > 0:
                start_idx = max(0, end_idx - word_span)
            
            # Extract the quote window
            quote_text = " ".join(quote_words[start_idx:end_idx])
        
        # Create a Text widget for formatted quote (supports bold)
        quote_text_widget = tk.Text(
            row_frame,
            font=("Arial", 20),
            bg=COLOR_BACKGROUND,
            fg=COLOR_TEXT,
            height=4,
            wrap=tk.WORD,
            relief=tk.FLAT,
            padx=5,
            pady=5
        )
        quote_text_widget.pack(side=tk.LEFT, padx=5, pady=5, fill=tk.BOTH, expand=True)
        
        # Insert text first
        quote_text_widget.insert("1.0", quote_text)
        
        # Apply bold formatting to search words (case-insensitive)
        # Normalize the quote for matching
        norm_quote = normalize_text(quote_text)
        quote_words_list = quote_text.split()
        norm_quote_words = [normalize_text(w) for w in quote_words_list]
        
        # For each search word, find matching words in the quote and bold them
        for search_word in search_words:
            norm_search_word = normalize_text(search_word)
            if not norm_search_word:
                continue
            
            # Find all words in the quote that match the search word
            for i, norm_quote_word in enumerate(norm_quote_words):
                if words_match(norm_search_word, norm_quote_word):
                    # Calculate the start and end positions of this word in the text widget
                    # Count characters before this word
                    char_count = sum(len(w) + 1 for w in quote_words_list[:i])  # +1 for space
                    start_pos = f"1.0+{char_count}c"
                    end_pos = f"1.0+{char_count + len(quote_words_list[i])}c"
                    quote_text_widget.tag_add("bold", start_pos, end_pos)
        
        quote_text_widget.tag_config("bold", font=("Arial", 20, "bold"))
        quote_text_widget.config(state=tk.DISABLED)  # Make read-only after formatting
    
    # Update scroll region
    results_canvas.update_idletasks()
    results_canvas.configure(scrollregion=results_canvas.bbox("all"))
    
    # Add export button at the bottom
    global export_button_frame
    if export_button_frame:
        # Clear existing buttons
        for widget in export_button_frame.winfo_children():
            widget.destroy()
        
        # Add export button
        export_btn = tk.Button(
            export_button_frame,
            text="📄 Creează raport",
            command=lambda: export_selected_to_docx(search_term),
            bg=COLOR_SHELF,
            fg="white",
            font=("Arial", 30, "bold"),
            relief=tk.RAISED,
            padx=20,
            pady=10,
            cursor="hand2"
        )
        export_btn.pack()


def open_match_image(match: dict):
    """Open the image for a search match."""
    volume_path_str = match.get("volume_path", "")
    image_filename = match.get("image", "")
    
    if not volume_path_str or not image_filename:
        messagebox.showerror("Error", "Image information not available.")
        return
    
    volume_path = Path(volume_path_str)
    image_path = find_image_path(volume_path, image_filename)
    
    # PDFs: open with PDF navigation (fallback to system viewer if deps missing)
    if image_path.suffix.lower() == ".pdf":
        open_pdf_with_navigation(image_path)
    else:
        open_image_with_navigation(image_path)


def export_selected_to_docx(search_term: str):
    """Export selected matches to a Word document."""
    global match_checkboxes
    
    # Get selected matches
    selected_matches = [match for match, checkbox_var in match_checkboxes if checkbox_var.get()]
    
    if not selected_matches:
        messagebox.showwarning("No selection", "Please select at least one result to export.")
        return
    
    try:
        # Get output directory
        output_dir = get_output_path()
        
        # Find or create default.docx template
        docx_file = None
        
        # Check if running as PyInstaller executable and look for bundled default.docx
        if getattr(sys, 'frozen', False) and hasattr(sys, '_MEIPASS'):
            bundled_docx = Path(sys._MEIPASS) / "default.docx"
            if bundled_docx.exists():
                docx_file = bundled_docx
        
        # If not found in bundle, check output_dir
        if docx_file is None or not docx_file.exists():
            docx_file = output_dir / "default.docx"
        
        # Load template if it exists, otherwise create new document
        if docx_file.exists():
            doc = Document(str(docx_file))
            # Clear existing content (keep styles)
            for para in list(doc.paragraphs):
                p = para._element
                p.getparent().remove(p)
        else:
            doc = Document()
        
        # Add title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run('Search results for "')
        search_run = title_para.add_run(search_term)
        search_run.italic = True
        title_para.add_run('"')
        doc.add_paragraph()
        
        # Add selected matches
        for match in selected_matches:
            # Get context: 5 lines above and 5 lines below the match
            if 'page_lines' in match and 'match_start_line_idx' in match:
                page_lines = match['page_lines']
                match_start = match.get('match_start_line_idx', 0)
                match_end = match.get('match_end_line_idx', match_start + 1)
                
                # Calculate context window: 5 lines above and 5 lines below
                context_start = max(0, match_start - 5)
                context_end = min(len(page_lines), match_end + 5)
                
                # Extract context lines
                context_lines = page_lines[context_start:context_end]
                
                # Add context lines to document
                for line in context_lines:
                    if line.strip():
                        para = doc.add_paragraph(line.strip())
                        try:
                            para.style = "2-Versuri-centru"
                        except KeyError:
                            pass
            else:
                # Fallback: use snippet if available, otherwise use matched_fragment
                if 'snippet' in match and match['snippet']:
                    for snippet_line in match['snippet']:
                        if snippet_line.strip():
                            para = doc.add_paragraph(snippet_line.strip())
                            try:
                                para.style = "2-Versuri-centru"
                            except KeyError:
                                pass
                elif 'matched_fragment' in match and match['matched_fragment']:
                    para = doc.add_paragraph(match['matched_fragment'].strip())
                    try:
                        para.style = "2-Versuri-centru"
                    except KeyError:
                        pass
            
            # Get title and code from match (exactly as displayed in table)
            title = match.get("title", "")
            code = match.get("code", "")
            volume_title = match.get("folder", "")
            page_num = match.get("page_num", "")
            
            # Capitalize only first letter of title
            if title:
                title = title[0].upper() + title[1:].lower() if len(title) > 0 else title
            
            # First line: TITLE (ITALIC) — (em dash) VOLUME TITLE, p. PAGE NUMBER (Arial 10)
            if title:
                title_para = doc.add_paragraph()
                # Apply style first
                try:
                    title_para.style = "3-Sursa text"
                except KeyError:
                    pass
                # Then add content with formatting
                title_run = title_para.add_run(title)
                title_run.italic = True
                title_para.add_run(f" — {volume_title}, p. {page_num}")
            else:
                # If no title, just show volume and page
                title_para = doc.add_paragraph(f"{volume_title}, p. {page_num}")
                # Apply style first
                try:
                    title_para.style = "3-Sursa text"
                except KeyError:
                    pass
            
            # Second line: THE CODE / CODES (Arial 10)
            # Always add code line, use "qqq" if no code found
            code_text = code if code else "qqq"
            code_para = doc.add_paragraph(code_text)
            # Apply style first
            try:
                code_para.style = "3-Sursa text"
            except KeyError:
                pass
            
            doc.add_paragraph()
        
        # Generate filename from search term
        suggested_filename = generate_filename_from_search(search_term)
        
        # Get last export directory or use default
        last_dir = get_last_export_dir()
        initial_dir = last_dir if last_dir else str(output_dir)
        
        # Open save dialog for user to choose location and filename
        save_path = filedialog.asksaveasfilename(
            title="Save Report",
            defaultextension=".docx",
            filetypes=[("Word documents", "*.docx"), ("All files", "*.*")],
            initialfile=suggested_filename,
            initialdir=initial_dir
        )
        
        if not save_path:
            # User cancelled
            return
        
        # Save document
        doc.save(save_path)
        
        # Remember the directory for next time (persists across app restarts)
        save_last_export_dir(str(Path(save_path).parent))
        
        # Show success message with DESCHIDE button
        success_dialog = tk.Toplevel(root)
        success_dialog.title("Success")
        success_dialog.configure(bg=COLOR_BACKGROUND)
        success_dialog.transient(root)
        success_dialog.grab_set()
        
        # Message label
        msg_text = f"Report exported:\n{save_path}\n\n{len(selected_matches)} result(s) exported."
        msg_label = tk.Label(
            success_dialog,
            text=msg_text,
            font=("Arial", 20),
            bg=COLOR_BACKGROUND,
            fg=COLOR_TEXT,
            justify=tk.LEFT,
            padx=20,
            pady=20
        )
        msg_label.pack()
        
        # Button frame
        button_frame = tk.Frame(success_dialog, bg=COLOR_BACKGROUND)
        button_frame.pack(pady=10)
        
        # DESCHIDE button
        deschide_btn = tk.Button(
            button_frame,
            text="📄 DESCHIDE",
            command=lambda: [open_document(Path(save_path)), success_dialog.destroy()],
            bg=COLOR_SHELF,
            fg="white",
            font=("Arial", 20, "bold"),
            relief=tk.RAISED,
            padx=15,
            pady=8,
            cursor="hand2"
        )
        deschide_btn.pack(side=tk.LEFT, padx=10)
        
        # OK button
        ok_btn = tk.Button(
            button_frame,
            text="OK",
            command=success_dialog.destroy,
            bg=COLOR_BOOK,
            fg=COLOR_TEXT,
            font=("Arial", 20, "bold"),
            relief=tk.RAISED,
            padx=15,
            pady=8,
            cursor="hand2"
        )
        ok_btn.pack(side=tk.LEFT, padx=10)
        
        # Center the dialog after widgets are packed
        success_dialog.update_idletasks()
        dialog_width = success_dialog.winfo_width()
        dialog_height = success_dialog.winfo_height()
        screen_width = success_dialog.winfo_screenwidth()
        screen_height = success_dialog.winfo_screenheight()
        x = (screen_width // 2) - (dialog_width // 2)
        y = (screen_height // 2) - (dialog_height // 2)
        success_dialog.geometry(f"+{x}+{y}")
        
        # Focus on the dialog
        success_dialog.focus_set()
    except Exception as e:
        # Show user-friendly error dialog
        messagebox.showerror(
            "Error",
            f"Failed to export report:\n{str(e)}"
        )


def on_generate():
    """Handle the Search button click."""
    global entry, word_span_var, word_order_var, button, status_var
    search_term = entry.get().strip()
    words = search_term.split()

    # Validate search term length (1-10 words)
    if not (1 <= len(words) <= 10):
        messagebox.showerror(
            "Invalid input",
            "Please enter between 1 and 10 words."
        )
        return

    # Check if root folder is selected
    if not selected_root_folder:
        messagebox.showerror(
            "Missing folder",
            "Please select a Biblioteca folder first."
        )
        return
    
    # Check if any volumes are selected
    if not selected_volumes:
        messagebox.showerror(
            "No selection",
            "Please select at least one folder to search."
        )
        return

    # Disable button during processing
    button.config(state="disabled")
    root.update()  # Update UI to show disabled state

    volume_count = len(selected_volumes)
    status_var.set(f"Volume selectate: {volume_count}. Pornesc căutarea...")
    search_start_ts = time.perf_counter()
    
    # Get word span value
    word_span_value = word_span_var.get()
    if word_span_value and word_span_value.isdigit() and int(word_span_value) > 0:
        word_span = int(word_span_value)
    else:
        word_span = 10  # Default to 10 words if invalid
    
    # Get word order setting (exact or random)
    order_value = word_order_var.get()
    exact_order = (order_value == "Exact")
    
    def search_worker():
        """Run search in background thread."""
        try:
            # Run search only (no document generation)
            matches = run_search_only(
                query_text=search_term,
                ocr_root=Path(selected_root_folder),
                word_span=word_span,
                threshold=DEFAULT_THRESHOLD,
                selected_volumes=selected_volumes,
                exact_order=exact_order
            )
            
            # Update GUI in main thread
            root.after(0, lambda: display_results_in_table(matches, search_term, word_span=word_span))
            elapsed = time.perf_counter() - search_start_ts
            root.after(0, lambda: status_var.set(f"Căutare finalizată: {volume_count} volume în {elapsed:.1f}s, {len(matches)} rezultate."))
            root.after(0, lambda: button.config(state="normal"))
            
        except Exception as e:
            # Show error in main thread
            root.after(0, lambda: messagebox.showerror(
                "Error",
                f"Failed to search:\n{str(e)}"
            ))
            root.after(0, lambda: status_var.set(f"Eșec căutare: {str(e)}"))
            root.after(0, lambda: button.config(state="normal"))
    
    # Start search in background thread
    search_thread = threading.Thread(target=search_worker, daemon=True)
    search_thread.start()


def _get_text_with_checkmark(text: str, add_checkmark: bool, tree_width: int = 600) -> str:
    """Get text with checkmark on the left before the title.
    tree_width: width of the tree column in pixels (default 600)"""
    # Remove existing checkmark if present (check both left and right)
    text = text.strip()
    if text.startswith("✓"):
        # Remove checkmark from the left
        text = text[1:].strip()
    elif "✓" in text:
        # Remove checkmark from the right (for backward compatibility)
        text = text.split("✓")[0].rstrip()
    
    if add_checkmark:
        # Put checkmark on the left before the text
        return "✓ " + text
    return text


def populate_library_tree(biblioteca_path: str):
    """Populate the library tree dynamically from folder structure.
    Marks folders containing ocr.txt as end folders (leaf nodes).
    All items are selected by default."""
    global library_tree
    library_tree.delete(*library_tree.get_children())
    root_path = Path(biblioteca_path)
    
    # Add root folder - selected by default (with checkmark aligned to right)
    root_text = _get_text_with_checkmark(f"📚 {root_path.name}", True)
    root_id = library_tree.insert("", "end", text=root_text, 
                                  values=(str(root_path), "folder"),
                                  tags=("folder", "selected"))
    
    # Recursively build tree
    _build_tree_recursive(root_path, root_id)
    
    # Expand root level
    library_tree.item(root_id, open=True)
    
    # Update selected volumes after populating tree
    update_selected_volumes()


def _build_tree_recursive(folder: Path, parent_item):
    """Recursively build tree structure, marking folders with OCR as end folders."""
    global library_tree
    # Check if this folder has OCR files
    has_ocr = folder_has_ocr(folder)
    
    # Get subfolders
    subfolders = sorted([p for p in folder.iterdir() if p.is_dir()])
    
    # If folder has OCR, mark it as end folder and don't show subfolders
    if has_ocr:
        # Mark parent as end folder (has OCR)
        current_tags = list(library_tree.item(parent_item, "tags"))
        if "end_folder" not in current_tags:
            library_tree.item(parent_item, tags=current_tags + ["end_folder"])
        return
    
    # If no OCR, show subfolders
    for subfolder in subfolders:
        # Determine icon and tags based on whether subfolder has OCR
        subfolder_has_ocr = folder_has_ocr(subfolder)
        
        # Get tree column width for proper alignment
        tree_width = library_tree.column("#0", "width") or 600
        
        if subfolder_has_ocr:
            # End folder (has OCR) - mark with special icon
            icon = "📗"
            tags = ("folder", "end_folder", "selected")
            text = _get_text_with_checkmark(f"{icon} {subfolder.name}", True, tree_width)  # Checkmark aligned to right
        else:
            # Regular folder (no OCR yet, may have subfolders)
            icon = "📁"
            tags = ("folder", "selected")
            text = _get_text_with_checkmark(f"{icon} {subfolder.name}", True, tree_width)  # Checkmark aligned to right
        
        subfolder_id = library_tree.insert(parent_item, "end", 
                                           text=text,
                                           values=(str(subfolder), "folder"),
                                           tags=tags)
        
        # Recursively process subfolders
        _build_tree_recursive(subfolder, subfolder_id)
    
    # Expand first level by default
    if parent_item:
        for child in library_tree.get_children(parent_item):
            library_tree.item(child, open=True)


def on_tree_check(event):
    """Handle checkbox clicks in the tree."""
    # Get the item that was clicked
    item = library_tree.selection()[0] if library_tree.selection() else None
    if not item:
        return
    
    # Toggle selection state
    current_tags = library_tree.item(item, "tags")
    if "selected" in current_tags:
        library_tree.item(item, tags=[tag for tag in current_tags if tag != "selected"])
    else:
        library_tree.item(item, tags=list(current_tags) + ["selected"])


def get_selected_paths():
    """Get all selected paths from the tree, handling parent-child relationships.
    Returns list of paths to search. If root folder is selected, returns just root.
    Otherwise returns selected folders (deduplicated)."""
    selected = []
    root_path = None
    
    # First pass: collect all selected items
    for item in library_tree.get_children():
        root_path, paths = _collect_selected(item, [])
        if root_path:
            # If root folder is selected, return it immediately (searches everything)
            return [str(root_path)]
        selected.extend(paths)
    
    # If no items selected, return empty list
    if not selected:
        return []
    
    # Deduplicate: if a parent is selected, remove its children
    # This handles cases like: Folder1 selected + SubFolder1 (child of Folder1) selected
    # We only need Folder1 since it includes SubFolder1
    selected_paths = []
    for path_str in selected:
        path = Path(path_str)
        # Check if this path is a child of any other selected path
        is_child = False
        for other_path_str in selected:
            if path_str != other_path_str:
                other_path = Path(other_path_str)
                try:
                    # Check if path is a subpath of other_path
                    path.relative_to(other_path)
                    is_child = True
                    break
                except ValueError:
                    pass
        
        if not is_child:
            selected_paths.append(path_str)
    
    return selected_paths


def _collect_selected(item, selected_list):
    """Recursively collect selected items. Returns (root_path, selected_paths).
    root_path is set if root folder itself is selected, otherwise None."""
    tags = library_tree.item(item, "tags")
    values = library_tree.item(item, "values")
    
    # Check if this item is selected
    if "selected" in tags and values:
        path_str = values[0]
        path = Path(path_str)
        
        # Check if this is the root item (first level child of tree root)
        parent = library_tree.parent(item)
        if parent == "":  # Root item
            return (Path(path_str), [])
        
        # Otherwise, add this path to the list
        selected_list.append(path_str)
    
    # Recursively check children
    for child in library_tree.get_children(item):
        root_path, child_paths = _collect_selected(child, [])
        if root_path:  # Root was selected in a child branch
            return (root_path, [])
        selected_list.extend(child_paths)
    
    return (None, selected_list)


def update_selected_volumes():
    """Update the global selected_volumes list with all selected end folders (with ocr.txt)."""
    global selected_volumes
    selected_volumes = []
    
    # Get all selected paths
    selected = get_selected_paths()
    
    # If root is selected, find all volumes recursively
    if selected_root_folder and len(selected) == 1 and Path(selected[0]) == Path(selected_root_folder):
        # Root selected - find all volumes
        _find_all_volumes(Path(selected_root_folder), selected_volumes)
    else:
        # Specific folders selected - extract only volumes (end folders with ocr.txt)
        for path_str in selected:
            path = Path(path_str)
            if folder_has_ocr(path):
                # This is a volume (end folder with OCR)
                selected_volumes.append(str(path))
            else:
                # Not a volume, search recursively for volumes within it
                _find_all_volumes(path, selected_volumes)


def _find_all_volumes(folder: Path, volumes_list: list):
    """Recursively find all volumes (folders with ocr.txt) within a folder."""
    if folder_has_ocr(folder):
        # This folder is a volume
        volumes_list.append(str(folder))
        return  # Don't search subfolders if this is an end folder
    
    # Search in subfolders
    for subfolder in sorted(p for p in folder.iterdir() if p.is_dir()):
        _find_all_volumes(subfolder, volumes_list)


def _select_item_recursive(item, select=True):
    """Recursively select or deselect an item and its children, updating text and colors."""
    global library_tree
    tags = list(library_tree.item(item, "tags"))
    current_text = library_tree.item(item, "text")
    
    if select:
        if "selected" not in tags:
            tags.append("selected")
    else:
        if "selected" in tags:
            tags.remove("selected")
    
    # Get tree column width for proper alignment
    tree_width = library_tree.column("#0", "width") or 600
    
    # Update text with checkmark aligned to the right
    new_text = _get_text_with_checkmark(current_text, select, tree_width)
    library_tree.item(item, text=new_text, tags=tags)
    
    # Recursively process children
    for child in library_tree.get_children(item):
        _select_item_recursive(child, select)


def _deselect_parent_if_needed(item):
    """Deselect parent when any child is deselected.
    Recursively propagates up the tree."""
    global library_tree
    parent = library_tree.parent(item)
    # If no parent (root level), stop
    if not parent or parent == "":
        return
    
    # Deselect the parent immediately (if it's selected)
    parent_tags = list(library_tree.item(parent, "tags"))
    if "selected" in parent_tags:
        parent_tags.remove("selected")
        tree_width = library_tree.column("#0", "width") or 600
        parent_text = library_tree.item(parent, "text")
        new_text = _get_text_with_checkmark(parent_text, False, tree_width)
        library_tree.item(parent, text=new_text, tags=parent_tags)
        # Recursively deselect parent's parent
        _deselect_parent_if_needed(parent)


def on_tree_click(event):
    """Handle clicks on tree items (toggle selection with cascading to children)."""
    global library_tree
    region = library_tree.identify_region(event.x, event.y)
    if region == "cell" or region == "tree":
        item = library_tree.identify_row(event.y)
        if item:
            tags = list(library_tree.item(item, "tags"))
            # Determine new selection state (toggle)
            should_select = "selected" not in tags
            # Apply selection state to this item and all its children recursively
            _select_item_recursive(item, should_select)
            # If item was deselected, check if parent should also be deselected
            if not should_select:
                _deselect_parent_if_needed(item)
            # Update the selected volumes list
            update_selected_volumes()


# Global variables for GUI (initialized in create_gui function)
root = None
COLOR_BACKGROUND = "#F5E6D3"
COLOR_SHELF = "#8B4513"
COLOR_BOOK = "#D4A574"
COLOR_TEXT = "#2C1810"
COLOR_ACCENT = "#A0522D"
main_canvas = None
main_scrollbar = None
scrollable_main_frame = None
library_tree = None
match_checkboxes = []
results_window = None
results_scrollable_frame = None
results_canvas = None
export_button_frame = None
results_title_label = None
entry = None
word_span_var = None
word_order_var = None
folder_label = None
select_folder_button = None
button = None
image_viewer_state = {
    "window": None,
    "images": [],           # for raster images
    "file_paths": [],       # generic list for pdfs/images
    "file_index": 0,        # current file position in file_paths
    "index": 0,             # image index (same as file_index for images)
    "page_index": 0,        # page inside current pdf
    "photo": None,
    "label": None,
    "title": None,
    "pdf_doc": None,
    "pdf_total": 0,
    "is_pdf": False,
    "resize_job": None,
}

def create_gui():
    """Create and initialize the GUI."""
    
    global root, COLOR_BACKGROUND, COLOR_SHELF, COLOR_BOOK, COLOR_TEXT, COLOR_ACCENT
    global main_canvas, main_scrollbar, scrollable_main_frame
    global library_tree, match_checkboxes
    global results_window, results_scrollable_frame, results_canvas
    global export_button_frame, results_title_label
    global entry, word_span_var, word_order_var
    global folder_label, select_folder_button, button
    
    # Create main window with library theme
    root = tk.Tk()
    root.title("Biblioteca - Căutare Documente")
    root.minsize(600, 450)
    root.configure(bg="#F5E6D3")  # Warm beige background like old books

    # Library-themed colors
    COLOR_BACKGROUND = "#F5E6D3"
    COLOR_SHELF = "#8B4513"  # Brown like wooden shelves
    COLOR_BOOK = "#D4A574"  # Tan like book covers
    COLOR_TEXT = "#2C1810"  # Dark brown text
    COLOR_ACCENT = "#A0522D"  # Sienna accent

    # Create scrollable frame for the entire window
    main_canvas = tk.Canvas(root, bg=COLOR_BACKGROUND, highlightthickness=0)
    main_scrollbar = ttk.Scrollbar(root, orient="vertical", command=main_canvas.yview)
    scrollable_main_frame = tk.Frame(main_canvas, bg=COLOR_BACKGROUND)

    scrollable_main_frame.bind(
        "<Configure>",
        lambda e: main_canvas.configure(scrollregion=main_canvas.bbox("all"))
    )

    main_canvas.create_window((0, 0), window=scrollable_main_frame, anchor="nw")
    main_canvas.configure(yscrollcommand=main_scrollbar.set)

    # Pack canvas and scrollbar
    main_canvas.pack(side="left", fill="both", expand=True)
    main_scrollbar.pack(side="right", fill="y")

    # Bind mousewheel to canvas
    def _on_main_mousewheel(event):
        main_canvas.yview_scroll(int(-1*(event.delta/120)), "units")
    main_canvas.bind_all("<MouseWheel>", _on_main_mousewheel)

    # Update canvas width when scrollable frame changes
    def configure_main_canvas_width(event):
        canvas_width = event.width
        main_canvas.itemconfig(main_canvas.find_all()[0], width=canvas_width)

    scrollable_main_frame.bind("<Configure>", lambda e: main_canvas.configure(scrollregion=main_canvas.bbox("all")))
    main_canvas.bind("<Configure>", configure_main_canvas_width)

    # Header frame with library title
    header_frame = tk.Frame(scrollable_main_frame, bg=COLOR_BACKGROUND, pady=15)
    header_frame.pack(fill=tk.X)

    title_label = tk.Label(
        header_frame,
        text="📚 BIBLIOTECA 📚",
        font=("Times", 50, "bold"),
        bg=COLOR_BACKGROUND,
        fg=COLOR_TEXT
    )
    title_label.pack()

    subtitle_label = tk.Label(
        header_frame,
        text="Selectează Biblioteca, Rafturi sau Volume pentru căutare",
        font=("Times", 26, "italic"),
        bg=COLOR_BACKGROUND,
        fg=COLOR_ACCENT
    )
    subtitle_label.pack(pady=(5, 0))

    # Biblioteca selection section
    biblioteca_frame = tk.Frame(scrollable_main_frame, bg=COLOR_BACKGROUND)
    biblioteca_frame.pack(padx=15, pady=10, fill=tk.X)

    select_folder_button = tk.Button(
        biblioteca_frame,
        text="📂 Selectează Biblioteca",
        command=on_select_folder,
        bg=COLOR_SHELF,
        fg="white",
        font=("Arial", 26, "bold"),
        relief=tk.RAISED,
        padx=15,
        pady=8,
        cursor="hand2"
    )
    select_folder_button.pack(side=tk.LEFT, padx=(0, 10))

    folder_label = tk.Label(
        biblioteca_frame,
        text="Nicio bibliotecă selectată",
        fg=COLOR_ACCENT,
        bg=COLOR_BACKGROUND,
        font=("Arial", 24, "italic"),
        anchor="w"
    )
    folder_label.pack(side=tk.LEFT, fill=tk.X, expand=True)

    # Library tree frame with scrollbars
    tree_frame = tk.Frame(scrollable_main_frame, bg=COLOR_BACKGROUND)
    tree_frame.pack(padx=15, pady=10, fill=tk.X)

    # Create scrollbars
    v_scrollbar = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL)
    h_scrollbar = ttk.Scrollbar(tree_frame, orient=tk.HORIZONTAL)

    # Create tree with library styling
    library_tree = ttk.Treeview(
        tree_frame,
        columns=("path", "type"),
        show="tree",
        yscrollcommand=v_scrollbar.set,
        xscrollcommand=h_scrollbar.set,
        selectmode="extended"
    )
    library_tree.column("#0", width=600, minwidth=200, stretch=True)
    library_tree.column("path", width=0, stretch=False)  # Hidden column
    library_tree.column("type", width=0, stretch=False)  # Hidden column

    # Configure treeview font and row height
    style = ttk.Style()
    style.configure("Treeview", font=("Arial", 24), rowheight=50)
    style.configure("Treeview.Heading", font=("Arial", 24, "bold"))

    # Configure tags for styling
    library_tree.tag_configure("folder", background="#F0E0C0", foreground=COLOR_TEXT)
    library_tree.tag_configure("end_folder", background="#E8D5B7", foreground=COLOR_TEXT)  # Folders with OCR
    library_tree.tag_configure("selected", background="#4A7C59", foreground="white")  # Green color for selected items

    # Pack scrollbars and tree
    v_scrollbar.config(command=library_tree.yview)
    h_scrollbar.config(command=library_tree.xview)
    v_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
    h_scrollbar.pack(side=tk.BOTTOM, fill=tk.X)
    library_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

    # Bind click events
    library_tree.bind("<Button-1>", on_tree_click)
    library_tree.bind("<Double-1>", lambda e: None)  # Prevent default double-click behavior

    # Selection control buttons
    selection_frame = tk.Frame(scrollable_main_frame, bg=COLOR_BACKGROUND)
    selection_frame.pack(padx=15, pady=5, fill=tk.X)

    def select_all_items():
        """Select all items in the tree."""
        for item in library_tree.get_children():
            _select_item_recursive(item, True)
        # Update the selected volumes list
        update_selected_volumes()

    def deselect_all_items():
        """Deselect all items in the tree."""
        for item in library_tree.get_children():
            _select_item_recursive(item, False)
        # Update the selected volumes list
        update_selected_volumes()

    select_all_btn = tk.Button(
        selection_frame,
        text="✓ Selectează Tot",
        command=select_all_items,
        bg=COLOR_BOOK,
        fg=COLOR_TEXT,
        font=("Arial", 22),
        relief=tk.RAISED,
        padx=10,
        pady=5,
        cursor="hand2"
    )
    select_all_btn.pack(side=tk.LEFT, padx=(0, 10))

    deselect_all_btn = tk.Button(
        selection_frame,
        text="✗ Deselectează Tot",
        command=deselect_all_items,
        bg=COLOR_BOOK,
        fg=COLOR_TEXT,
        font=("Arial", 22),
        relief=tk.RAISED,
        padx=10,
        pady=5,
        cursor="hand2"
    )
    deselect_all_btn.pack(side=tk.LEFT)

    # Instructions label
    # instructions_label = tk.Label(
    #     root,
    #     text="💡 Click pe foldere pentru a le selecta/deselecta. Folderele cu 📗 conțin ocr.txt",
    #     font=("Arial", 9),
    #     bg=COLOR_BACKGROUND,
    #     fg=COLOR_ACCENT,
    #     pady=5
    # )
    # instructions_label.pack()

    # Search section
    search_frame = tk.Frame(scrollable_main_frame, bg=COLOR_BACKGROUND)
    search_frame.pack(padx=15, pady=10, fill=tk.X)

    tk.Label(
        search_frame,
        text="Căutare:",
        font=("Arial", 26, "bold"),
        bg=COLOR_BACKGROUND,
        fg=COLOR_TEXT
    ).pack(anchor="w", pady=(0, 5))

    entry = tk.Entry(
        search_frame,
        font=("Arial", 26),
        relief=tk.SUNKEN,
        bd=2
    )
    entry.pack(fill=tk.X, pady=5)

    # Word span section
    word_span_frame = tk.Frame(scrollable_main_frame, bg=COLOR_BACKGROUND)
    word_span_frame.pack(padx=15, pady=5, fill=tk.X)

    tk.Label(
        word_span_frame,
        text="Span cuvinte:",
        font=("Arial", 24),
        bg=COLOR_BACKGROUND,
        fg=COLOR_TEXT
    ).pack(side=tk.LEFT, padx=(0, 10))

    # Create dropdown for word span options
    word_span_var = tk.StringVar(value="10")  # Default to 10 words
    word_span_options = ["5", "10", "15", "20"]

    word_span_dropdown = ttk.Combobox(
        word_span_frame,
        textvariable=word_span_var,
        values=word_span_options,
        state="readonly",
        font=("Arial", 36)
    )
    word_span_dropdown.pack(side=tk.LEFT)

    # Word order section
    word_order_frame = tk.Frame(scrollable_main_frame, bg=COLOR_BACKGROUND)
    word_order_frame.pack(padx=15, pady=5, fill=tk.X)

    tk.Label(
        word_order_frame,
        text="Ordinea cuvintelor:",
        font=("Arial", 24),
        bg=COLOR_BACKGROUND,
        fg=COLOR_TEXT
    ).pack(side=tk.LEFT, padx=(0, 10))

    # Create dropdown for word order options
    word_order_var = tk.StringVar(value="Exact")  # Default to exact order
    word_order_options = ["Exact", "Aleatorie"]

    word_order_dropdown = ttk.Combobox(
        word_order_frame,
        textvariable=word_order_var,
        values=word_order_options,
        state="readonly",
        font=("Arial", 36)
    )
    word_order_dropdown.pack(side=tk.LEFT)

    # Generate button
    button = tk.Button(
        scrollable_main_frame,
        text="🔍 Caută",
        command=on_generate,
        bg=COLOR_SHELF,
        fg="white",
        font=("Arial", 30, "bold"),
        relief=tk.RAISED,
        padx=20,
        pady=10,
        cursor="hand2"
    )
    button.pack(pady=15)

    # Status label under the search button (shows volumes and duration)
    global status_var
    status_var = tk.StringVar(value="")
    status_label = tk.Label(
        scrollable_main_frame,
        textvariable=status_var,
        font=("Arial", 18, "italic"),
        bg=COLOR_BACKGROUND,
        fg=COLOR_ACCENT,
        justify=tk.LEFT
    )
    status_label.pack(pady=(0, 10), anchor="w", padx=20)

    # Global variable to store results window
    results_window = None
    results_scrollable_frame = None
    results_canvas = None
    export_button_frame = None
    results_title_label = None
    match_checkboxes = []  # List of (match, checkbox_var) tuples for export

def main():
    create_gui()
    if root is not None:
        root.mainloop()


if __name__ == "__main__":
    # Required on Windows/PyInstaller to prevent child processes from re-running GUI
    multiprocessing.freeze_support()
    main()

